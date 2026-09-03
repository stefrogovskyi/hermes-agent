import json
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

data = json.load(open('/opt/hermes/profiles/archie/final_article_data.json'))

doc = docx.Document()

# Add Title as H1
h1 = doc.add_heading(data['title'], level=1)

# Add Meta-Title and Meta-Description as italic 9pt paragraphs
p_meta1 = doc.add_paragraph()
run_m1 = p_meta1.add_run(f"Meta Title: {data['meta_title']}")
run_m1.italic = True
run_m1.font.size = Pt(9)

p_meta2 = doc.add_paragraph()
run_m2 = p_meta2.add_run(f"Meta Description: {data['meta_description']}")
run_m2.italic = True
run_m2.font.size = Pt(9)

doc.add_paragraph() # spacing

# Parse body text by sections/paragraphs
body = data['body']
lines = body.split('\n')

for line in lines:
    line_str = line.strip()
    if not line_str:
        continue
    if line_str.startswith('## '):
        doc.add_heading(line_str[3:], level=2)
    elif line_str.startswith('# '):
        doc.add_heading(line_str[2:], level=1)
    else:
        doc.add_paragraph(line_str)

docx_path = "/opt/hermes/profiles/archie/How_Practical_Data_Reporting_Fixes_Supply_Chain_Delays.docx"
doc.save(docx_path)
print("Saved DOCX to:", docx_path)
