import json
import subprocess
import sys
from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# Set standard margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title (H1)
title_text = "Prorated Billing Strategies for International Logistics Startups"
h1 = doc.add_heading(title_text, level=1)

# Meta Title (Italic 9pt)
meta_title_text = "Meta Title: Prorated Billing for International Logistics Startups"
p_mt = doc.add_paragraph()
r_mt = p_mt.add_run(meta_title_text)
r_mt.font.size = Pt(9)
r_mt.font.italic = True

# Meta Description (Italic 9pt)
meta_desc_text = "Meta Description: Learn how prorated billing for logistics helps international startups adjust subscriptions, build trust, and calculate partial charges accurately."
p_md = doc.add_paragraph()
r_md = p_md.add_run(meta_desc_text)
r_md.font.size = Pt(9)
r_md.font.italic = True

doc.add_paragraph() # spacing

# Content structure: list of tuples (type, text)
# type: 'p' = normal paragraph, 'h2' = heading level 2
content = [
    ('p', "Approximately 90% of startups fail when entering the competitive global market. Poor customer experience, such as delayed deliveries or damaged goods, frequently drives clients to alternative providers. Modern cargo shippers expect flexible financial terms alongside reliable physical transit. Offering adaptable billing models provides growing logistics companies with a strong foundation for customer retention."),
    
    ('h2', "Understanding the International Logistics Proration Process"),
    ('p', "Clients often pay upfront for long-distance transportation or tracking tools. Questions arise when service stops early or shipping requirements shift mid-period. Proration solves these mid-cycle adjustments by calculating exact service usage and issuing proportional refunds or fee adjustments."),
    ('p', "For instance, a client pays upfront for a full month of cargo tracking but cancels after twenty days. Proration determines the exact monetary value of those unused ten days and returns the balance. Without fair calculations, clients perceive unrefunded charges as unjust, which harms brand reputation and prompts word-of-mouth criticism."),
    ('p', "Mid-contract service changes also require financial adjustments. When a company upgrades to accelerated delivery halfway through a billing window, prorated billing for logistics guarantees that the higher rate applies only to the remaining active days. Explaining these terms clearly through email, customer service chat, or website notices builds long-term client confidence."),
    
    ('h2', "Calculating Proration with Precision"),
    ('p', "Logistics firms manage SaaS subscription adjustments and partial-period cargo tracking billing by applying a straightforward formula:"),
    ('p', "(Total Subscription Fee ÷ Total Billing Days) × Active Usage Days = Final Prorated Charge"),
    ('p', "Suppose a customer uses a shipment monitoring plan for twenty consecutive days during a 30-day month before unsubscribing, with the monthly service costing $200. Because the client retained active access across 20 days, the proration period equals 20 days."),
    ('p', "Dividing the $200 total cost by 30 days and multiplying the result by 20 days gives the required charge:"),
    ('p', "(200 ÷ 30) × 20 = $133.30"),
    ('p', "The logistics startup bills $133.30 instead of $200 and refunds the remaining balance. This process accurately reflects the service rendered for a fraction of a billing timeframe."),
    
    ('h2', "Structuring Partial-Month and Partial-Year Systems"),
    ('p', "Subscription-based logistics companies implement either partial-month or partial-year proration frameworks depending on contract structure:"),
    ('p', "1. Partial-Month Proration: Calculate active usage days as a fraction of the full monthly cycle, then multiply by the total cost to determine the precise payment amount."),
    ('p', "2. Partial-Year Proration: Measure how many active days a client used the service during the full year rather than a monthly cycle. Startups use this calculation to determine amounts owed for unused annual services."),
    ('p', "Adopting clear proration models gives clients confidence that they pay only for what they consume. Transparent accounting improves customer conversion, boosts freight forwarder billing retention, and stabilizes deferred revenue for long-term growth.")
]

for block_type, text in content:
    if block_type == 'h2':
        doc.add_heading(text, level=2)
    else:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)

filename = "/opt/hermes/profiles/archie/Prorated_Billing_Strategies_for_International_Logistics_Startups.docx"
doc.save(filename)
print(f"Docx file created successfully: {filename}")

# Upload to Google Drive
drive_folder_id = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
cmd = [
    "python3", "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py",
    "drive", "upload",
    "--parent", drive_folder_id,
    "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    filename
]

res = subprocess.run(cmd, capture_output=True, text=True)
if res.returncode != 0:
    print(f"Drive upload error: {res.stderr}")
    sys.exit(1)

print("Drive Upload Output:")
print(res.stdout)

upload_data = json.loads(res.stdout)
file_id = upload_data.get("id")
web_link = upload_data.get("webViewLink")

print(f"SUCCESS: File ID = {file_id}")
print(f"webViewLink = {web_link}")

with open("/opt/hermes/profiles/archie/upload_result_292.json", "w") as f:
    json.dump(upload_data, f, indent=2)
