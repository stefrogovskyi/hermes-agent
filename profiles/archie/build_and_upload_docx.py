import json
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    with open('/opt/hermes/profiles/archie/final_rewrite.json', 'r') as f:
        data = json.load(f)

    title = data["title"]
    meta_title = data["meta_title"]
    meta_desc = data["meta_description"]
    body = data["body_markdown"]

    doc = Document()

    # H1 - Title
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run(title)
    run_h1.font.size = Pt(20)
    run_h1.font.bold = True

    # Meta Title (Italic 9pt)
    p_meta1 = doc.add_paragraph()
    r_meta1 = p_meta1.add_run(f"Meta Title: {meta_title}")
    r_meta1.font.size = Pt(9)
    r_meta1.font.italic = True

    # Meta Description (Italic 9pt)
    p_meta2 = doc.add_paragraph()
    r_meta2 = p_meta2.add_run(f"Meta Description: {meta_desc}")
    r_meta2.font.size = Pt(9)
    r_meta2.font.italic = True

    doc.add_paragraph() # Spacer

    # Body Markdown parsing
    lines = body.split('\n')
    in_list = False

    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
        
        if line_str.startswith('# '):
            h = doc.add_heading(line_str[2:], level=1)
        elif line_str.startswith('## '):
            h = doc.add_heading(line_str[3:], level=2)
        elif line_str.startswith('### '):
            h = doc.add_heading(line_str[4:], level=3)
        elif line_str.startswith('* ') or line_str.startswith('- '):
            item_text = line_str[2:]
            p = doc.add_paragraph(item_text, style='List Bullet')
        else:
            p = doc.add_paragraph(line_str)

    docx_path = "/opt/hermes/profiles/archie/SeaRates_Updates_Week_41_2024.docx"
    doc.save(docx_path)
    print("Saved DOCX to", docx_path)

    # Upload to Google Drive
    parent_id = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
    cmd = [
        "python3",
        "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py",
        "drive", "upload", docx_path,
        "--parent", parent_id,
        "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    
    print("Uploading to Drive...")
    res = subprocess.run(cmd, capture_output=True, text=True)
    print("Upload result stdout:", res.stdout)
    print("Upload result stderr:", res.stderr)

    if res.returncode == 0:
        try:
            drive_res = json.loads(res.stdout)
            file_id = drive_res.get("id")
            web_view_link = drive_res.get("webViewLink", f"https://drive.google.com/file/d/{file_id}/view")
            print("Successfully uploaded. webViewLink:", web_view_link)
            with open('/opt/hermes/profiles/archie/upload_result.json', 'w') as f_out:
                json.dump({"file_id": file_id, "webViewLink": web_view_link, "drive_res": drive_res}, f_out, indent=2)
        except Exception as e:
            print("Error parsing upload result json:", e)
    else:
        print("Upload failed!")

if __name__ == "__main__":
    main()
