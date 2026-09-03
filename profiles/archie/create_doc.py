import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(filename):
    doc = docx.Document()

    # Title as H1
    h1 = doc.add_heading('SeaRates Load Calculator Video Guide', level=1)

    # Meta Title in italic 9pt
    p_meta_t = doc.add_paragraph()
    r_meta_t = p_meta_t.add_run('Meta Title: SeaRates Load Calculator Guide: Cargo Space Optimization')
    r_meta_t.italic = True
    r_meta_t.font.size = Pt(9)

    # Meta Description in italic 9pt
    p_meta_d = doc.add_paragraph()
    r_meta_d = p_meta_d.add_run('Meta Description: Watch our freight loading video guide to see how the SeaRates Load Calculator helps plan container shipments with speed and accuracy.')
    r_meta_d.italic = True
    r_meta_d.font.size = Pt(9)

    doc.add_paragraph() # Blank line separator

    # Body paragraphs
    body_paragraphs = [
        "Fitting cargo into a truck or container without wasting space takes precise planning. To help shippers avoid calculation errors and load freight faster, SeaRates offers its Load Calculator, a smart stuffing tool for digital shipment planning.",
        "In this video guide, Account Manager Kateryna Kernesh walks through the tool step by step. You will learn how this container load calculator simplifies cargo space optimization, keeps cargo safe, and makes shipments more profitable.",
        "Watch the manual on our YouTube channel at https://youtu.be/Ez9DNjhs5WU?si=Bk4ZAJs-jDzNGKJp to start planning your shipments with clarity.",
        "If you have feedback or want to request a customized Load Calculator for your business, contact our team at it.sales@searates.com or submit a request for an IT Quote."
    ]

    for p_text in body_paragraphs:
        doc.add_paragraph(p_text)

    doc.save(filename)
    print(f"Saved {filename}")

if __name__ == '__main__':
    create_docx('/opt/hermes/profiles/archie/SeaRates_Load_Calculator_Video_Guide.docx')
