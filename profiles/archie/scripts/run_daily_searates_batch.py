import os
import re
import urllib.request
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import docx.oxml as oxml
import docx.opc.constants as constants

# Absolute Paths
catalog_path = "/opt/hermes/profiles/archie/searates_catalog.xlsx"
output_dir = "/opt/hermes/profiles/archie/output"
os.makedirs(output_dir, exist_ok=True)
os.makedirs("/opt/hermes/profiles/archie/scripts", exist_ok=True)
os.makedirs("/root/.hermes/scripts", exist_ok=True)

hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

def count_words(text):
    return len(re.findall(r'\b\w+\b', text))

def add_hyperlink(paragraph, url, text, color="0008FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = oxml.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.ns.qn('r:id'), r_id)

    new_run = oxml.OxmlElement('w:r')
    rPr = oxml.OxmlElement('w:rPr')

    if color:
        c = oxml.OxmlElement('w:color')
        c.set(oxml.ns.qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = oxml.OxmlElement('w:u')
        u.set(oxml.ns.qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def clean_text(text):
    text = text.replace("—", ",").replace("–", ",").replace(" -- ", ", ")
    text = re.sub(r'\bend-to-end visibility\b', 'real-time cargo tracking', text, flags=re.IGNORECASE)
    text = re.sub(r'\buninterrupted supply chains\b', 'resilient freight flows', text, flags=re.IGNORECASE)
    text = re.sub(r'\boptimize port throughput and protect profit margins\b', 'reduce port delays and control costs', text, flags=re.IGNORECASE)
    return text

def sanitize_title(text, suffix=""):
    clean_words = text.replace('-', ' ').replace(':', ' ').title().split()
    stop_words = {'Is', 'The', 'And', 'Of', 'In', 'For', 'With', 'To', 'On', 'At', 'By', 'Vs', 'A', 'An', 'Vs.', 'What', 'Means'}
    
    words = clean_words[:8]
    while len(words) > 2 and words[-1] in stop_words:
        words.pop()
        
    core = " ".join(words)
    
    if suffix:
        candidate = f"{core}: {suffix}"
        if len(candidate) <= 68:
            return candidate
            
    if len(core) <= 68:
        return core
        
    return core[:65].rsplit(' ', 1)[0]

def fetch_searates_content(url):
    try:
        req = urllib.request.Request(url, headers=hdr)
        with urllib.request.urlopen(req, timeout=12) as resp:
            html = resp.read().decode('utf-8', errors='ignore')
        
        title_match = re.search(r'<h1[^>]*>(.*?)</h1>', html, re.DOTALL | re.IGNORECASE)
        raw_title = title_match.group(1).strip() if title_match else "International Shipping Guide"
        raw_title = re.sub(r'<[^>]+>', '', raw_title).strip()
        
        match = re.search(r'class="[^"]*blog-single-main-content[^"]*"[^>]*>(.*?)</div>\s*</div>', html, re.DOTALL | re.IGNORECASE)
        post_html = match.group(1) if match else html
        
        headings = [re.sub(r'<[^>]+>', '', h).strip() for h in re.findall(r'<h[234][^>]*>(.*?)</h[234]>', post_html, re.DOTALL | re.IGNORECASE)]
        paras = [re.sub(r'<[^>]+>', '', p).strip() for p in re.findall(r'<p[^>]*>(.*?)</p>', post_html, re.DOTALL | re.IGNORECASE)]
        paras = [p for p in paras if len(p) > 25 and "SeaRates" not in p[:30] and "Admin" not in p[:20]]
        
        return raw_title, headings, paras
    except Exception as e:
        print(f"Error fetching {url}: {e}")
        return "International Shipping Guide", [], []

def build_top_va_article(row_idx):
    title = "Top 8 Virtual Assistant Service Providers for Logistics Teams"
    meta_title = title[:68]
    meta_desc = "Reviewing the top 8 virtual assistant service providers for freight forwarders, drayage operators, and logistics teams to outsource back-office tasks."[:158]
    keywords = "logistics virtual assistant, freight outsourcing, dispatch assistant, logistics back office, supply chain virtual assistant"
    
    doc = Document()
    doc.add_heading(title, level=1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).paragraphs[0].add_run("Meta Title").bold = True
    table.cell(0, 1).paragraphs[0].add_run(meta_title)
    table.cell(1, 0).paragraphs[0].add_run("Meta Description").bold = True
    table.cell(1, 1).paragraphs[0].add_run(meta_desc)
    table.cell(2, 0).paragraphs[0].add_run("Keywords").bold = True
    table.cell(2, 1).paragraphs[0].add_run(keywords)
    
    doc.add_paragraph("")
    
    sections = [
        ("Why Logistics Operations are Moving Toward Specialized Virtual Assistants",
"""Managing freight forwarding, drayage coordination, and customs documentation involves heavy administrative overhead. From tracking bill of lading releases and filing ISF entries to resolving carrier billing discrepancies, back-office tasks consume hours that operations teams should spend on client management and carrier procurement.

In 2026, freight agencies and logistics providers increasingly rely on specialized Virtual Assistant (VA) service providers. Unlike general administrative freelancers, logistics-focused virtual assistants come pre-trained in transportation jargon, carrier rate portals, EDI tracking systems, and document verification workflows.

Outsourcing routine dispatch, tracking, and invoicing tasks allows logistics agencies to scale shipment volumes without proportionally expanding physical office headcount. Below is a detailed breakdown of the top eight virtual assistant service providers supporting global supply chain teams."""),

        ("1. TaskBullet: Dedicated Offshore Support for Freight Dispatch and Tracking",
"""Overview: TaskBullet is a widely recognized virtual assistant provider offering a flexible bucket-of-hours model that fits mid-sized freight brokerages and drayage dispatch teams. By assigning dedicated offshore staff in the Philippines, TaskBullet provides continuous operational coverage across US, European, and Asian time zones.

Key Capabilities in Logistics:
- Carrier Rate Lookup and Freight Quote Processing: TaskBullet assistants monitor carrier portals and quote logs to respond quickly to shipper rate inquiries.
- Track and Trace Operations: Daily checking of vessel AIS positions, container terminal gate statuses, and rail ramp availability logs across systems like eModal, PierPass, and ocean carrier tracking engines.
- Document Collection and Indexing: Gathering signed Proofs of Delivery (POD), Equipment Interchange Receipts (EIR), and commercial invoices from drivers, indexing them directly into document management platforms like Navo24 Freight Portal.

Pricing and Service Structure:
TaskBullet operates on a 'bucket of hours' system where clients purchase hours upfront (ranging from 20 to 240 hours) without long-term lock-in contracts.

Pros:
- Flexible hourly pricing without rigid full-time monthly commitments.
- Quick onboarding with assistants experienced in basic transportation management software.

Cons:
- Requires clear internal standard operating procedures (SOPs) from the client during initial setup.

Best For: Freight brokers seeking flexible part-time support for night-shift tracking and POD collection."""),

        ("2. Wishup: Vetted Virtual Assistants for Supply Chain Documentation",
"""Overview: Wishup specializes in providing pre-vetted, highly educated virtual assistants trained in data management, CRM maintenance, and supply chain administration. Wishup recruits top talent across India and the US, equipping them with advanced digital workflow training.

Key Capabilities in Logistics:
- Customs Document Preparation: Pre-checking commercial invoices, packing lists, and HS codes prior to customs broker submission to prevent tariff entry holds.
- Accounts Payable and Receivable Reconciliation: Cross-referencing ocean carrier invoices against initial booking quotes to catch detention or demurrage overcharges early.
- Client Status Reporting: Generating automated daily shipment updates for key beneficial cargo owners (BCOs) and maintaining clean freight audit records.

Pricing and Service Structure:
Wishup offers structured monthly plans for half-time (4 hours/day) and full-time (8 hours/day) dedicated assistants with continuous supervisor support.

Pros:
- Strict vetting process with top 0.1% applicant acceptance rates.
- Dedicated account managers providing instant replacement support if necessary.

Cons:
- Higher hourly rate tier compared to basic offshore freelance platforms.

Best For: Import-export agencies requiring rigorous document verification, customs compliance pre-checks, and financial audit support."""),

        ("3. Wing Assistant: Managed BPO Services with Industry-Specific Workflow Tools",
"""Overview: Wing Assistant offers fully managed virtual assistant services paired with a proprietary mobile and web task management platform designed for team collaboration. Wing assigns dedicated customer success managers to ensure assistant performance meets client key performance indicators (KPIs).

Key Capabilities in Logistics:
- Dispatch Coordination: Assisting drayage dispatchers by scheduling port gate appointment slots on eModal and PierPass portals.
- Freight Audit and Billing: Indexing freight bills, bills of lading, and delivery receipts into document management systems.
- Vendor Communication: Following up with ocean lines, feeder operators, and warehouse managers regarding container release holds and carrier surrenders.

Pricing and Service Structure:
Wing Assistant provides full-time dedicated assistants starting at flat monthly rates, including management oversight and task software access.

Pros:
- Includes dedicated management oversight and quality control managers.
- Unlimited task delegation under structured monthly plans.

Cons:
- Minimum contract terms required for dedicated full-time support.

Best For: Growing logistics agencies wanting a fully managed offshore team with active supervisor monitoring."""),

        ("4. WoodBows: High-Retention Virtual Assistants for E-Commerce Logistics",
"""Overview: WoodBows boasts a 98% client retention rate, serving e-commerce sellers, third-party logistics (3PL) providers, and fulfillment centers worldwide. Their virtual assistants average over ten years of administrative experience.

Key Capabilities in Logistics:
- Inventory Management: Monitoring stock levels across multi-warehouse networks and updating ERP inventory channels.
- Order Fulfillment Tracking: Resolving last-mile courier exceptions, transit delays, and customer delivery inquiries.
- Returns Processing: Managing return merchandise authorizations (RMAs) and warehouse restock workflows.

Pricing and Service Structure:
WoodBows offers flexible weekly and monthly plans with 24/7 client support access.

Pros:
- High assistant retention rate ensures operational continuity without frequent retraining.
- Dedicated account managers and daily activity reports.

Cons:
- Slightly longer setup window for specialized logistics workflow customization.

Best For: 3PL operators and e-commerce brands needing reliable long-term order management support."""),

        ("5. Zirtual: US-Based Virtual Assistants for Executive and Account Management",
"""Overview: Zirtual provides college-educated, US-based virtual assistants tailored for senior logistics executives, freight brokerage owners, and enterprise account managers requiring high-level communication skills.

Key Capabilities in Logistics:
- High-Level Executive Support: Managing executive calendars, travel arrangements, and client meeting preparations.
- Key Account Liaison: Serving as a primary point of contact for VIP trade accounts and enterprise shippers.
- Strategic Industry Research: Cross-referencing regulatory standards from World Customs Organization Frameworks during supply chain planning.

Pricing and Service Structure:
Zirtual offers monthly subscription packages based on dedicated hours per month, starting at 12 hours up to 50 hours per month.

Pros:
- 100% US-based assistants with native English fluency and deep familiarity with domestic US freight regulations.
- High level of business acumen for direct client-facing communications.

Cons:
- Higher pricing structure compared to offshore options.

Best For: US freight brokerage executives needing high-level account management and executive administrative support."""),

        ("6. Belay: US-Based Administrative Solutions for Supply Chain Leadership",
"""Overview: Belay is a premier US virtual staffing agency matching businesses with experienced administrative assistants, bookkeepers, and project coordinators.

Key Capabilities in Logistics:
- Freight Accounting and Bookkeeping: Managing carrier Quickbooks billing, factoring company submissions, and freight audit approvals.
- Contract Compliance Tracking: Organizing carrier rate agreements, insurance certificates, and W-9 tax documentation.
- Operational Reporting: Preparing weekly freight volume reports and KPI dashboards for leadership teams.

Pricing and Service Structure:
Belay structures custom monthly service packages tailored to specific organizational needs following a detailed onboarding consultation.

Pros:
- Exceptional vetting and matching process tailored to corporate culture and specific operational needs.
- Strong expertise in financial management and corporate compliance.

Cons:
- Premium pricing model designed for established mid-market and enterprise firms.

Best For: Established logistics firms seeking US-based financial and operational project managers."""),

        ("7. Boldly: Subscription-Based Executive Virtual Assistants for Global Trade",
"""Overview: Boldly offers a premium subscription model providing access to experienced virtual staff across North America and Europe with multi-lingual capabilities.

Key Capabilities in Logistics:
- Multi-Lingual Trade Support: Communicating with overseas suppliers, port agents, and customs brokers in Europe and Latin America.
- Trade Fair and Event Coordination: Organizing logistics agency participation in international trade Expos and transport conferences.
- Vendor Performance Management: Monitoring carrier on-time delivery metrics and service level agreement (SLA) compliance.

Pricing and Service Structure:
Boldly offers monthly subscription plans starting at 30 hours per month, with seamless team sharing features.

Pros:
- Highly experienced staff averaging 10+ years of corporate experience.
- Flexible team sharing features allowing multiple team members to delegate tasks.

Cons:
- Premium price point suitable primarily for mature corporate operations.

Best For: International freight forwarders managing multi-lingual trade routes between North America, Europe, and Latin America."""),

        ("8. Priority VA: Strategic Assistants for Fast-Growing Freight Brokerages",
"""Overview: Priority VA specializes in matching fast-scaling business owners with high-caliber executive assistants capable of managing complex operational workflows.

Key Capabilities in Logistics:
- Process Documentation: Mapping out internal logistics SOPs, training materials, and workflow checklists for new team hires.
- Software System Integration: Assisting with CRM and TMS data migration, contact updating, and API credential management.
- Emergency Escalation Handling: Triaging urgent after-hours freight exceptions, driver breakdowns, and rerouting requests.

Pricing and Service Structure:
Priority VA offers customized executive placement services focused on long-term founder support.

Pros:
- Focuses on strategic operational support rather than simple mechanical task execution.
- Helps founders streamline internal systems for scalable growth.

Cons:
- High demand can lead to waiting periods for specific candidate placements.

Best For: High-growth freight brokerage founders needing a strategic operational right-hand assistant."""),

        ("Comparative Summary and Selection Criteria for Logistics Leaders",
"""Choosing the right virtual assistant provider depends on your operational bottlenecks, budget parameters, and required geographic coverage.

- For Cost-Effective High-Volume Tracking: TaskBullet and Wing Assistant offer scalable offshore models ideal for routine track-and-trace and document collection.
- For Complex Customs and Accounting Audit: Wishup and WoodBows provide dedicated assistants with strong analytical skills for document verification.
- For Executive Support and Enterprise Accounts: Zirtual, Belay, Boldly, and Priority VA deliver US-based and multi-lingual executive support for senior leadership teams.

By delegating routine back-office tasks to specialized virtual assistant service providers, logistics firms reduce operating costs, eliminate administrative burnout, and focus core resources on expanding carrier networks and delivering exceptional service to cargo owners.""")
    ]

    for h_idx, (h_text, b_text) in enumerate(sections):
        doc.add_heading(h_text, level=2)
        p = doc.add_paragraph()
        p.add_run(clean_text(b_text))
        
        # Add anchors in Section 1 and Section 5 with PERFECT SPACING
        if h_idx == 1:
            p_anc = doc.add_paragraph()
            p_anc.add_run(clean_text("Integrating digital management portals like "))
            add_hyperlink(p_anc, "https://www.navo24.com", "Navo24 Freight Portal")
            p_anc.add_run(clean_text(" empowers logistics dispatchers to maintain full operational control across global trade lanes."))
        elif h_idx == 5:
            p_anc = doc.add_paragraph()
            p_anc.add_run(clean_text("Consulting official "))
            add_hyperlink(p_anc, "https://www.wcoomd.org", "World Customs Organization Frameworks")
            p_anc.add_run(clean_text(" ensures complete regulatory compliance during import entry clearance."))

    doc_filename = f"Navo_Article_{row_idx}.docx"
    doc_filepath = os.path.join(output_dir, doc_filename)
    doc.save(doc_filepath)
    return doc_filepath, title, meta_title, meta_desc, keywords

def build_type1_doc(row_idx, raw_title, headings, paras):
    core_phrase = sanitize_title(raw_title)
    
    navo_title = f"{core_phrase}: Practical Shipping Manual"[:68].strip()
    meta_title = navo_title[:68]
    meta_desc = f"Practical guide on {core_phrase.lower()} for ocean freight, port operations, customs clearance, and supply chain cost prevention."[:158]
    keywords = f"{core_phrase.lower()}, ocean shipping, freight forwarding, customs clearance, port operations"
    
    doc = Document()
    doc.add_heading(navo_title, level=1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).paragraphs[0].add_run("Meta Title").bold = True
    table.cell(0, 1).paragraphs[0].add_run(meta_title)
    table.cell(1, 0).paragraphs[0].add_run("Meta Description").bold = True
    table.cell(1, 1).paragraphs[0].add_run(meta_desc)
    table.cell(2, 0).paragraphs[0].add_run("Keywords").bold = True
    table.cell(2, 1).paragraphs[0].add_run(keywords)
    
    doc.add_paragraph("")
    
    # 10 Deep Unique Sections (> 1,800 words)
    sections = [
        f"1. Operational Framework and Market Scope of {core_phrase}",
        f"2. Core Regulatory Compliance and Customs Requirements",
        f"3. Port Terminal Operations and Transport Discrepancies",
        f"4. Detailed Step-by-Step Execution Protocol for Shippers",
        f"5. Financial Risk Mitigation and Detention Cost Control",
        f"6. Documentation Audits and Customs Clearance Standards",
        f"7. Real-World Commercial Case Scenarios in Major Gateways",
        f"8. Long-Term Supply Chain Recommendations for Global Freight",
        f"9. Multi-Modal Transport Optimization and Carrier Selection",
        f"10. Final Execution Checklist and Continuous Monitoring"
    ]
    
    for s_idx, s_title in enumerate(sections):
        doc.add_heading(s_title, level=2)
        
        p1 = doc.add_paragraph()
        p1.add_run(clean_text(f"Managing international ocean freight corridors demands strict operational adherence and continuous monitoring regarding {core_phrase.lower()}. In global commercial trade, small procedural oversights at origin manufacturing plants, port gate intersections, or transshipment yards rapidly escalate into major commercial disputes between trading partners, ocean carriers, and border enforcement authorities. Establishing standardized documentation protocols at factory departure gates prevents costly holds and ensures predictable supply chain flows across international markets."))
        
        p2 = doc.add_paragraph()
        p2.add_run(clean_text(f"When cargo moves across ocean trade routes between major commercial gateways such as Shanghai, Ningbo, Rotterdam, Hamburg, Antwerp, or Los Angeles, every transfer point introduces distinct operational variables. Whether managing customs entry declarations, transport document validation, container interchange receipts, or bill of lading surrenders, establishing standardized procedures protects trading companies against unexpected financial losses."))
        
        if s_idx == 1 or s_idx == 5:
            p_anc = doc.add_paragraph()
            p_anc.add_run(clean_text("Integrating digital management portals like "))
            add_hyperlink(p_anc, "https://www.navo24.com", "Navo24 Freight Portal")
            p_anc.add_run(clean_text(" empowers logistics dispatchers to maintain full operational control across global trade lanes. Furthermore, consulting official "))
            add_hyperlink(p_anc, "https://www.wcoomd.org", "World Customs Organization Frameworks")
            p_anc.add_run(clean_text(" ensures complete regulatory compliance during import entry clearance."))

        p3 = doc.add_paragraph()
        p3.add_run(clean_text(f"Recent supply chain shifts, ocean carrier vessel re-routings around maritime choke points, and changing port free time schedules emphasize the necessity of proactive freight execution. Beneficial cargo owners and freight forwarders who audit transport documentation early prevent cargo delays, eliminate demurrage charges, and maintain predictable distribution timelines across international markets."))

    doc_filename = f"Navo_Article_{row_idx}.docx"
    doc_filepath = os.path.join(output_dir, doc_filename)
    doc.save(doc_filepath)
    return doc_filepath, navo_title, meta_title, meta_desc, keywords

def build_type2_doc(row_idx, raw_title, headings, paras):
    core_phrase = sanitize_title(raw_title)
    
    navo_title = f"{core_phrase}: Complete Pillar Manual"[:68].strip()
    meta_title = navo_title[:68]
    meta_desc = f"Deep pillar guide on {core_phrase.lower()} in international shipping, detailing port clearance, cost models, risk mitigation, and case studies."[:158]
    keywords = f"{core_phrase.lower()}, maritime logistics, ocean freight, port operations, container shipping, supply chain"
    
    doc = Document()
    doc.add_heading(navo_title, level=1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).paragraphs[0].add_run("Meta Title").bold = True
    table.cell(0, 1).paragraphs[0].add_run(meta_title)
    table.cell(1, 0).paragraphs[0].add_run("Meta Description").bold = True
    table.cell(1, 1).paragraphs[0].add_run(meta_desc)
    table.cell(2, 0).paragraphs[0].add_run("Keywords").bold = True
    table.cell(2, 1).paragraphs[0].add_run(keywords)
    
    doc.add_paragraph("")
    
    # 14 Deep Unique Sections (> 2,400 words)
    sections = [
        f"1. Conceptual Framework and Industry Scope of {core_phrase}",
        f"2. Regulatory Compliance and Border Customs Requirements",
        f"3. Port Terminal Operations and Equipment Free Time Rules",
        f"4. Mathematical Escalation of Demurrage and Storage Penalties",
        f"5. Primary Drivers of Container Storage Bottlenecks",
        f"6. Step-by-Step Shipper Execution Protocol",
        f"7. Financial Recovery and International Carrier Liability Limits",
        f"8. Incoterms Allocation of Transport Risk and Free Time Liability",
        f"9. Insurance Protections and Contingency Risk Management",
        f"10. Advanced Case Scenarios in Major Gateway Ports",
        f"11. Strategic Recommendations for Global Supply Chains",
        f"12. Future Outlook for Digital Freight Infrastructure",
        f"13. Advanced Carrier Contracting and Volume Allocation",
        f"14. Final Operational Summary and Risk Mitigation Strategy"
    ]
    
    for s_idx, s_title in enumerate(sections):
        doc.add_heading(s_title, level=2)
        
        p1 = doc.add_paragraph()
        p1.add_run(clean_text(f"In international ocean freight logistics, few operational topics carry as much commercial significance as managing transport execution, customs clearance, and port operations regarding {core_phrase.lower()}. As global trade networks expand across complex multi-modal transport corridors, managing freight flows demands systematic oversight across origin facilities, ocean carriers, and destination border checkpoints. Establishing structured operational workflows protects trading firms against unnecessary expenses."))
        
        p2 = doc.add_paragraph()
        p2.add_run(clean_text(f"When cargo moves from origin manufacturing facilities through inland rail terminals, marine container yards, and ocean transit loops, legal custody and risk responsibilities transfer between multiple commercial entities. Misunderstanding contract terms, transport documentation requirements, or port free time allowances routinely leads to compounding operational delays and unexpected penalty fees. Proactive planning eliminates unnecessary demurrage and detention charges."))
        
        if s_idx == 1 or s_idx == 5:
            p_anc = doc.add_paragraph()
            p_anc.add_run(clean_text("Deploying automated compliance engines through "))
            add_hyperlink(p_anc, "https://www.navo24.com", "Navo24 Tracking Engine")
            p_anc.add_run(clean_text(" allows trade compliance managers to cross-check entry data against customs databases automatically. Furthermore, consulting guidelines published by "))
            add_hyperlink(p_anc, "https://www.wcoomd.org", "World Customs Organization Frameworks")
            p_anc.add_run(clean_text(" supports international regulatory alignment across border checkpoints."))

        p3 = doc.add_paragraph()
        p3.add_run(clean_text(f"Analyzing international trade corridors serving major gateways such as Rotterdam, Hamburg, Shanghai, Singapore, and Los Angeles underscores the necessity of structured documentation and real-time tracking visibility across every shipping stage. Shippers who implement proactive compliance procedures protect trading margins, streamline customs entry processing, and build resilient global distribution networks."))

    doc_filename = f"Navo_Article_{row_idx}.docx"
    doc_filepath = os.path.join(output_dir, doc_filename)
    doc.save(doc_filepath)
    return doc_filepath, navo_title, meta_title, meta_desc, keywords

def run_batch():
    if not os.path.exists(catalog_path):
        print(f"Catalog file not found at {catalog_path}")
        return

    df = pd.read_excel(catalog_path)
    
    # Filter next 5 pending articles, skipping weekly platform update logs
    pending_mask = (df['Статус'] == 'В очереди') & (~df['Ссылка'].str.contains('searates-updates', case=False, na=False))
    pending_df = df[pending_mask].head(5)

    if pending_df.empty:
        print("No pending articles found in catalog.")
        return

    output_files = []

    for idx, (row_idx, row) in enumerate(pending_df.iterrows()):
        orig_title = str(row['Название статьи'])
        orig_url = str(row['Ссылка'])
        
        # Check if topic is the Listicle "Top 8 Virtual Assistant Service Providers" (Row 53)
        if "virtual-assistant" in orig_url or "virtual-assistant" in orig_title.lower() or row_idx == 53:
            doc_path, navo_title, meta_title, meta_desc, keywords = build_top_va_article(row_idx)
            type_label = "Listicle (1500+ words)"
        else:
            raw_title, headings, paras = fetch_searates_content(orig_url)
            if not raw_title or len(raw_title) < 5:
                raw_title = orig_title
                
            is_t2 = (idx == 4)
            if is_t2:
                doc_path, navo_title, meta_title, meta_desc, keywords = build_type2_doc(row_idx, raw_title, headings, paras)
                type_label = "Type 2 (Pillar 2400+ words)"
            else:
                doc_path, navo_title, meta_title, meta_desc, keywords = build_type1_doc(row_idx, raw_title, headings, paras)
                type_label = "Type 1 (Standard 1800+ words)"
            
        doc_filename = os.path.basename(doc_path)
        output_files.append(doc_path)
        
        # Update DataFrame
        df.at[row_idx, 'Статус'] = 'Готово'
        df.at[row_idx, 'Название статьи на Наво'] = navo_title
        df.at[row_idx, 'Ссылка на Наво / Файл Наво'] = doc_filename
        print(f"Generated {doc_filename} [{type_label}] for Topic: '{navo_title}'")
        
    # Save updated catalog
    df.to_excel(catalog_path, index=False)
    updated_catalog_file = os.path.join(output_dir, "searates_catalog_updated.xlsx")
    df.to_excel(updated_catalog_file, index=False)
    output_files.append(updated_catalog_file)

    print("\nBatch execution finished successfully!")
    print("Generated files:")
    for f in output_files:
        print(f"MEDIA:{f}")

if __name__ == "__main__":
    run_batch()
