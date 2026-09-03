import json
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('/opt/hermes/profiles/archie/verified_article.json', 'r', encoding='utf-8') as f:
    article = json.load(f)

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title (H1)
h1 = doc.add_heading(article['title'], level=1)
h1.style.font.name = 'Arial'
h1.style.font.size = Pt(20)
h1.style.font.bold = True
h1.style.font.color.rgb = RGBColor(0, 51, 102)

# Meta Title (Italic 9pt)
p_meta_t = doc.add_paragraph()
run_mt_label = p_meta_t.add_run("Meta Title: ")
run_mt_label.font.name = 'Arial'
run_mt_label.font.size = Pt(9)
run_mt_label.font.italic = True
run_mt_label.font.bold = True
run_mt_val = p_meta_t.add_run(article['meta_title'])
run_mt_val.font.name = 'Arial'
run_mt_val.font.size = Pt(9)
run_mt_val.font.italic = True

# Meta Description (Italic 9pt)
p_meta_d = doc.add_paragraph()
run_md_label = p_meta_d.add_run("Meta Description: ")
run_md_label.font.name = 'Arial'
run_md_label.font.size = Pt(9)
run_md_label.font.italic = True
run_md_label.font.bold = True
run_md_val = p_meta_d.add_run(article['meta_description'])
run_md_val.font.name = 'Arial'
run_md_val.font.size = Pt(9)
run_md_val.font.italic = True

# Spacer
doc.add_paragraph()

# Body parsing
body = article['body']
lines = body.split('\n')

for line in lines:
    line_str = line.strip()
    if not line_str:
        continue
    
    if line_str.startswith('### '):
        # Section Heading (H2)
        heading_text = line_str.replace('### ', '').strip()
        h2 = doc.add_heading(heading_text, level=2)
        h2.style.font.name = 'Arial'
        h2.style.font.size = Pt(14)
        h2.style.font.bold = True
        h2.style.font.color.rgb = RGBColor(51, 51, 51)
    elif line_str.startswith('## '):
        heading_text = line_str.replace('## ', '').strip()
        h2 = doc.add_heading(heading_text, level=2)
        h2.style.font.name = 'Arial'
        h2.style.font.size = Pt(14)
        h2.style.font.bold = True
        h2.style.font.color.rgb = RGBColor(51, 51, 51)
    elif line_str.startswith('- '):
        bullet_text = line_str[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bullet_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    else:
        p = doc.add_paragraph()
        run = p.add_run(line_str)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

docx_path = "/opt/hermes/profiles/archie/How_to_Use_the_Demurrage_Storage_Calculator.docx"
doc.save(docx_path)
print(f"DOCX created successfully at {docx_path}")
