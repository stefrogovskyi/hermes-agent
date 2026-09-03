import os
import re
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import docx.oxml as oxml
import docx.opc.constants as constants

output_dir = "/opt/hermes/profiles/archie/output"
os.makedirs(output_dir, exist_ok=True)

def count_words(text):
    return len(re.findall(r'\b\w+\b', text))

def add_hyperlink(paragraph, url, text, color="0008FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = oxml.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.ns.qn('r:id'), r_id)

    new_run = oxml.OxmlElement('w:r')
    rPr = oxml.OxmlElement('w:rPr')

    if color:
        c = oxml.OxmlElement('w:color')
        c.set(oxml.ns.qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = oxml.OxmlElement('w:u')
        u.set(oxml.ns.qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def clean_text(text):
    text = text.replace("—", ",").replace("–", ",").replace(" -- ", ", ")
    text = re.sub(r'\bend-to-end visibility\b', 'real-time cargo tracking', text, flags=re.IGNORECASE)
    text = re.sub(r'\buninterrupted supply chains\b', 'resilient freight flows', text, flags=re.IGNORECASE)
    text = re.sub(r'\boptimize port throughput and protect profit margins\b', 'reduce port delays and control costs', text, flags=re.IGNORECASE)
    text = re.sub(r'\bvital role\b', 'key operational role', text, flags=re.IGNORECASE)
    text = re.sub(r'\bdelve into\b', 'examine', text, flags=re.IGNORECASE)
    text = re.sub(r'\bgame-changer\b', 'major operational advantage', text, flags=re.IGNORECASE)
    text = re.sub(r'\bin today\'s world\b', 'in modern freight operations', text, flags=re.IGNORECASE)
    return text

def build_ai_chat_article():
    navo_title = "How AI Assistant Chat Streamlines Freight Forwarding Workflows"
    meta_title = "AI Assistant Chat for Freight Workflows: Operations Guide"
    meta_desc = "Learn how AI assistant chat tools eliminate repetitive freight queries, automate cargo tracking, and streamline logistics operations for forwarders."
    keywords = "AI assistant chat, freight workflows, logistics automation, ocean freight tracking, TMS integration, freight forwarding AI"

    # 10 Deep, Rich sections (~190-210 words per section = ~1,550+ words total)
    sections = [
        ("1. Operational Overview: Addressing Repetitive Back-Office Inquiries in Freight Forwarding",
         """In global freight forwarding operations, up to 70 percent of daily customer communication consists of repetitive operational inquiries. Freight dispatchers, customer service representatives, and logistics coordinators spend hours manually answering routine questions regarding vessel positions, container terminal gate releases, freight quote statuses, and customs document clearance holds.

When operations staff spend most of their workday copying container numbers, cross-referencing carrier tracking portals, and drafting repetitive email updates, strategic tasks like carrier rate negotiations, route optimization, and client account growth take a backseat. Operating under manual communication workflows limits shipment scalability and increases administrative error rates across international trade lanes.

Integrating specialized conversational AI assistants into transportation management systems (TMS) transforms how freight agencies handle routine inquiries. By connecting AI assistants directly to live ocean carrier APIs, AIS vessel tracking feeds, and internal ERP databases, logistics firms automate routine communications while maintaining complete operational accuracy.

Proactive communication models allow logistics teams to handle higher shipment volumes without proportionally expanding customer support headcount, protecting operating margins during volatile market cycles."""),

        ("2. The Anatomy of Routine Freight Communication: Rate Quotes, AIS Tracking, and Gate Status",
         """Analyzing back-office inquiry logs reveals consistent patterns across customer communication channels. The vast majority of shipper inquiries fall into four primary categories: container location inquiries, freight quote requests, customs document status updates, and invoicing clarifications.

- Container Location Tracking: Shippers demand real-time location updates regarding ocean containers moving across transpacific and Asia-Europe trade lanes.
- Freight Rate Inquiries: Beneficial Cargo Owners (BCOs) frequently request spot market ocean freight quotes, drayage transport estimates, and bunker adjustment factor (BAF) breakdowns.
- Customs and Port Release Hold Checks: Importers check whether customs entry summary declarations have cleared border agencies or whether terminal demurrage clocks have started.
- Invoicing and Document Verification: Freight accountants resolve billing discrepancies, verify proof of delivery (POD) receipts, and confirm bill of lading releases.

Automating these four communication categories with AI conversational assistants frees up operations staff to focus on high-value logistical troubleshooting, risk mitigation, and active customer relationship management."""),

        ("3. Eliminating Manual Data Re-Entry Across Carrier Portals and TMS Platforms",
         """A major inefficiency in traditional freight forwarding is data fragmentation. Dispatchers constantly switch between multiple ocean line portals, port terminal appointment platforms like eModal or PierPass, customs broker feeds, and internal TMS databases to find single shipment details.

Manual data entry across disparate systems creates operational bottlenecks and introduces transcription errors. When a dispatcher manually types a 10-digit container booking number into three different tracking portals to answer a single phone call, labor efficiency drops significantly across the entire office.

AI assistant chat interfaces eliminate manual data switching by aggregating multi-carrier tracking feeds into a single conversational interface. Shippers and internal account managers can query shipment statuses via natural language chat messages, receiving instant, verified data pulled directly from live carrier systems.

Centralizing tracking data improves data integrity, speeds up customer response times, and eliminates the frustration of searching through fragmented software platforms."""),

        ("4. AI Assistant Chat Workflows: Automating Instant Shipper Responses Without Sacrificing Accuracy",
         """Deploying AI chat assistants in freight forwarding requires strict operational safeguards to ensure data accuracy. Unlike generic AI chatbots, logistics-focused assistants operate under structured business logic tied to verified transportation databases.

When a shipper submits a query like 'What is the estimated arrival time for container MSKU9876543 at the Port of Rotterdam?', the AI assistant executes a multi-step verification sequence:

First, it validates the container number against active TMS booking records. Second, it queries live AIS satellite tracking data and ocean carrier schedules to verify vessel position. Third, it formats a concise, human-readable update specifying estimated berth arrival, terminal gate status, and customs clearance holds.

Providing instant, accurate responses 24/7 improves client satisfaction while reducing phone call volumes for busy dispatch desks during peak shipping seasons."""),

        ("5. Automated Document Verification: Pre-Checking Bills of Lading, Commercial Invoices, and HS Codes",
         """Beyond answering tracking queries, advanced AI assistants streamline cargo documentation workflows. Incomplete commercial invoices, incorrect packing list weights, and misdeclared Harmonized System (HS) codes cause severe customs holds and port demurrage charges.

AI assistants pre-screen uploaded shipping documents prior to customs entry submission. By scanning digital PDFs of commercial invoices, packing lists, and marine bills of lading, the assistant cross-references data fields to detect discrepancies early.

For example, if the total carton count on a packing list differs from the commercial invoice total, the AI assistant flags the error to the dispatcher before filing entry summary documents with border authorities, preventing costly customs holds and exam fees.

Automated pre-auditing of shipping documentation protects importers against regulatory fines and ensures audit readiness across global supply chains."""),

        ("6. Managing Logistics Exceptions and Delay Alerts: Proactive Exception Handling via Conversational AI",
         """While routine shipments move smoothly, severe weather, port labor strikes, and blank sailings cause unexpected transport exceptions. Traditional reactive communication means shippers discover delays only after missing scheduled delivery windows.

Conversational AI assistants enable proactive exception management. When AIS satellite feeds indicate a vessel is delayed by port congestion or weather disruptions, the AI assistant automatically identifies affected container bookings and drafts personalized client notification summaries.

Instead of dispatchers frantically drafting hundreds of individual delay emails, the AI assistant alerts shippers proactively, providing updated estimated arrival times, alternative drayage options, and revised port appointment recommendations.

Proactive communication turns potential customer service crises into demonstrations of operational competence and reliability."""),

        ("7. Step-by-Step Implementation Protocol for Logistics Agencies and Freight Brokerages",
         """To successfully deploy AI assistant chat workflows without disrupting daily freight operations, logistics managers should follow a structured four-step rollout protocol:

Step 1: Audit Communication Channels. Review customer support tickets and email logs to identify top recurring inquiry topics across specific trade routes and shipping corridors.

Step 2: Connect Core Data Sources. Integrate the AI chat interface with internal TMS platforms, ocean carrier tracking APIs, and customs broker EDI channels.

Step 3: Establish Fallback Guardrails. Configure automatic human handoff rules so complex shipping exceptions, insurance claims, or billing disputes transfer immediately to senior dispatchers.

Step 4: Conduct Supervised Pilot Testing. Run the AI assistant internally with customer service teams for two weeks before opening direct client-facing chat channels.

Following a disciplined rollout process ensures seamless integration, protects client relationships, and builds team confidence in automated AI workflows."""),

        ("8. Real-World Case Scenario: Streamlining Dispatch and Customer Support in Busy Ocean Corridors",
         """Consider a mid-sized freight forwarding agency managing 2,000 FEU monthly across Asia-US West Coast trade routes. During peak shipping season, the customer service team received over 400 daily tracking inquiries via email and phone calls.

After deploying an integrated AI assistant chat tool connected to their TMS, routine tracking inquiry emails dropped by 65 percent within thirty days. Shippers retrieved instant container location updates, gate availability notices, and vessel schedules directly through a self-service portal.

Customer support representatives redirected saved working hours toward securing drayage trucking capacity, negotiating spot ocean rates, and expanding service offerings for core beneficial cargo owners.

The agency scaled monthly shipment volumes by 30 percent without adding administrative support staff, significantly boosting net operating margins across core shipping lanes."""),

        ("9. Data Security, Regulatory Compliance, and System Integration Safeguards",
         """Integrating AI conversational tools into supply chain infrastructure requires strict data security measures. Transport documentation contains sensitive commercial information, including commodity valuations, shipper identities, and proprietary rate structures.

Logistics AI platforms must enforce enterprise-grade security protocols, including end-to-end data encryption, role-based access controls, and compliance with global data privacy standards. Restricting database access ensures shippers view only their authorized container bookings.

Furthermore, maintaining detailed audit logs of all AI-generated communications ensures operational transparency and provides clear documentation records during post-clearance customs audits.

Prioritizing data security safeguards client trust and ensures compliance with international trade regulatory standards across all global operating locations."""),

        ("10. Long-Term Strategic Value: Scaling Freight Volumes and Operating Margins Without Increasing Headcount",
         """Adopting conversational AI assistants in freight forwarding is not merely an administrative convenience; it is a strategic driver of operating margin growth. As ocean freight markets fluctuate, forwarders who keep back-office operating costs low maintain strong competitive advantages.

Automating up to 70 percent of routine customer inquiries enables freight agencies to double shipment handling capacity without doubling administrative staff headcount. Dispatchers transition from manual data processors into strategic supply chain advisors.

Combining automated tracking visibility, instant rate quotes, document verification, and proactive exception alerts empowers modern logistics agencies to deliver superior customer service while protecting operating profitability across global trade lanes.

Forward-thinking logistics providers who embrace intelligent automation establish resilient, scalable business models built for long-term commercial success in competitive transport markets.""")
    ]

    doc = Document()
    doc.add_heading(navo_title, level=1)

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).paragraphs[0].add_run("Meta Title").bold = True
    table.cell(0, 1).paragraphs[0].add_run(meta_title)
    table.cell(1, 0).paragraphs[0].add_run("Meta Description").bold = True
    table.cell(1, 1).paragraphs[0].add_run(meta_desc)
    table.cell(2, 0).paragraphs[0].add_run("Keywords").bold = True
    table.cell(2, 1).paragraphs[0].add_run(keywords)

    doc.add_paragraph("")

    for s_idx, (s_title, s_text) in enumerate(sections):
        doc.add_heading(s_title, level=2)
        p = doc.add_paragraph()
        p.add_run(clean_text(s_text))

        # Hyperlink Anchors with EXACT Spacing
        if s_idx == 2:
            p_anc = doc.add_paragraph()
            p_anc.add_run(clean_text("Integrating digital logistics portals like "))
            add_hyperlink(p_anc, "https://www.navo24.com", "Navo24 Freight Portal")
            p_anc.add_run(clean_text(" empowers forwarders to centralize ocean tracking data and streamline client communications."))
        elif s_idx == 8:
            p_anc = doc.add_paragraph()
            p_anc.add_run(clean_text("Adhering to international standard regulations published by the "))
            add_hyperlink(p_anc, "https://www.wcoomd.org", "World Customs Organization Frameworks")
            p_anc.add_run(clean_text(" ensures complete compliance during automated data transmission."))

    # Stage 7 Audit Gate
    full_text = " ".join([p.text for p in doc.paragraphs if p.text.strip()])
    wc = count_words(full_text)
    has_emdash = "—" in full_text or "--" in full_text
    has_images = len(doc.inline_shapes) > 0

    audit_passed = (wc >= 1500) and (not has_emdash) and (not has_images)
    print(f"==========================================")
    print(f"[Stage 7 Audit Gate] Verdict: {'PASS' if audit_passed else 'FAIL'}")
    print(f"Word Count: {wc} words (Min required: 1500)")
    print(f"Em-dashes Present: {has_emdash}")
    print(f"Images Present: {has_images}")
    print(f"==========================================")

    doc_filepath = os.path.join(output_dir, "Navo_Article_AI_Chat_Freight.docx")
    doc.save(doc_filepath)
    print(f"File saved to: {doc_filepath}")
    print(f"MEDIA:{doc_filepath}")

if __name__ == "__main__":
    build_ai_chat_article()
