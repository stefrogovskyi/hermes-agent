import docx
from docx import Document
from docx.shared import Pt

title = "Perfect Schedules, Broken Shipments: What Really Fails"
meta_title = "Why Accurate Ocean Freight Schedules Still Fail"
meta_desc = "Accurate ETAs don't guarantee smooth shipments. See why ocean freight plans still break, and what separates stable schedules from fragile ones."

doc = Document()
doc.add_heading(title, level=1)

p = doc.add_paragraph()
run = p.add_run(f"Meta Title: {meta_title}")
run.italic = True
run.font.size = Pt(9)

p2 = doc.add_paragraph()
run2 = p2.add_run(f"Meta Description: {meta_desc}")
run2.italic = True
run2.font.size = Pt(9)

doc.add_paragraph("")

sections = [
    (None, [
        "A shipping schedule is a promise written in pencil, dressed up to look like ink.",
        "You have the vessel schedule. The ETA. Transit time. A handful of routing options to pick from. On paper, the plan is airtight. And yet containers still show up on the wrong day, warehouse slots go unused, and someone on your team is scrambling to rebuild a delivery sequence that was supposed to run itself. If you built your plan around a carrier's published timetable and nothing else, this probably sounds familiar.",
        "Planning decisions usually boil down to three inputs: transit time, expected arrival, and which routes are even available. That works fine when a service performs exactly the way it's advertised. Right now, that's not a safe bet to make.",
    ]),
    ("The slot nobody warned you about", [
        "Say the vessel lands two days later than the schedule promised. Two days doesn't sound like much until you watch what it does downstream: the warehouse slot you booked is gone, unloading gets pushed to whenever there's room, and every delivery scheduled after that one shifts with it.",
        "That damage was already locked in before the ship ever left port, the moment the plan assumed the schedule would hold exactly as printed.",
        "From there the costs stack quietly. Trucking gets rebooked. Cargo sits in storage longer than budgeted. Buffer costs appear on invoices nobody planned for, all traced back to a delay that looked minor on the day it happened.",
    ]),
    ("Same ETA, different animal", [
        "Leaning on the headline ETA as your only planning input is common, and it's a habit worth breaking. The primary ETA tells you when a vessel is supposed to show up. It says nothing about how firmly that date is likely to hold.",
        "Put two carrier services side by side and they can carry the identical ETA on the exact same lane, yet behave nothing alike once cargo is actually moving. One holds to a tight window sailing after sailing, while the other swings between early and late for no obvious reason. The booking page shows the same figure for both; their track records on past voyages don't match up nearly as closely.",
    ]),
    ("Myths that quietly wreck the plan", [
        "A few assumptions keep showing up in planning rooms, and none of them survive contact with real operations for long: that schedules hold still once published, that services on the same lane perform roughly the same, that a shorter transit time automatically means a better outcome, and that small deviations don't really matter in the end. Each one sounds reasonable on its own. Stack enough of them together and the plan stops reflecting how the lane actually behaves.",
    ]),
    ("Where the plan actually comes apart", [
        "And the route looked fine too, chosen for the fastest transit time on offer, the schedule clean and confident on the screen. That schedule doesn't include the service's actual track record on past voyages.",
        "Most of the damage is done before the cargo ever leaves the dock, quietly, during the routing decision that looked perfectly reasonable at the time. A service that shifts its schedule often starts drifting before the ship even sails. One with wide swings in arrival times turns your arrival window into a guess. And when port rotations change mid-route, the coordination you built downstream starts to unravel with it.",
        "Underneath every published schedule sits a pattern most planners never look at: how often the service revises itself before departure, how far actual arrivals stray from the plan, how steady the port rotation stays voyage after voyage, and how fast the service bounces back once something goes wrong. Nothing on a standard schedule view shows any of that, though all of it decides whether the shipment stays on plan.",
        "A schedule is more than a date on a screen. A useful one behaves like a running update, something you check against reality and use to predict freight deadlines with some confidence. A fixed line printed once and trusted blindly can't do that.",
        "Strip away the context, the stability record, the deviation history, the pattern of recovery after things go wrong, and an ETA turns into a fairly weak signal. Shippers, carriers, managers, the people actually moving freight day to day, tend to read a clear date as a reliable one. That distinction between a clear date and a genuinely reliable one is easy to miss until a delay exposes the gap.",
    ]),
]

for heading, paras in sections:
    if heading:
        doc.add_heading(heading, level=2)
    for para in paras:
        doc.add_paragraph(para)

# Table section
doc.add_heading("Typical planning vs. performance-based planning", level=2)
table_data = [
    ("", "Typical planning approach", "Performance-based approach"),
    ("Primary focus", "The ETA and how long transit takes", "Live ETA tracking combined with a service's track record"),
    ("Decision question", "Which option gets there quickest?", "Which option holds to its schedule?"),
    ("View of the schedule", "Something fixed and settled", "Something that updates as conditions change"),
    ("Risk assessment", "Judged from the ETA by itself", "Judged from ongoing updates and past behavior"),
    ("Service comparison", "Appears identical at a glance", "Ranked by how it has actually performed"),
    ("Reaction to delays", "Handled once they've already happened", "Planned for before they occur"),
    ("Planning outcome", "Prone to breaking down", "Holds up under pressure"),
]
table = doc.add_table(rows=len(table_data), cols=3)
table.style = 'Table Grid'
for i, row in enumerate(table_data):
    for j, cell_text in enumerate(row):
        cell = table.cell(i, j)
        cell.text = cell_text
        if i == 0:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph("")

sections2 = [
    ("What \"reliable\" is even supposed to mean", [
        "A reliable service doesn't reveal itself in a single ETA. It shows up across voyage after voyage, in how tightly the arrivals cluster around what was promised.",
        "The strongest schedule isn't necessarily the quickest one on the board. It's steady enough that downstream coordination doesn't need constant rebuilding, consistent enough that you're not repeatedly reworking the plan, and predictable enough that freight cost and risk stay inside a range you can actually manage. That combination is what turns raw schedule data into something you can build a plan on.",
    ]),
    ("Three ways to actually check the data", [
        "Reading service behavior takes more than staring at one schedule table. It means checking that same booking through a few different lenses.",
        "Start with a route-based search: what gets you from A to B. This is the wide-angle view, carrier options laid out side by side, complete with transit durations, arrival estimates, and other routing choices worth weighing. It's the fastest way to map the market and shortlist what's worth a closer look.",
        "Next comes vessel or service-level search. You leave the wide route view behind and zero in on a single service. Here you're checking how often its schedule changes before departure, how consistent its arrivals have been across past voyages, how stable the whole service has stayed over time. This is the step that separates two options that look identical at booking but behave completely differently once cargo is moving.",
        "Then there's the port view, which asks a different question entirely: what's happening at this specific port right now. It tracks upcoming arrivals and departures, gives a live read on how services calling that terminal are doing right now, and pulls in real-time updates across the carriers that use it. Used this way, schedules stop being read in isolation. You can spot port congestion signals early, compare overlapping services hitting the same port, and update the plan around what's actually happening on the ground, replacing whatever last week's plan assumed.",
    ]),
    ("Closing thought", [
        "Being able to see how routes, vessels, and terminals actually behave is what turns scattered schedule data into something you can build a plan on. Skip that context and even a perfectly accurate schedule won't rescue a plan that assumed too much.",
        "Send your question to it.sales@searates.com. SeaRates works from schedules built to actually perform, giving your planning something real to stand on.",
    ]),
]

for heading, paras in sections2:
    doc.add_heading(heading, level=2)
    for para in paras:
        doc.add_paragraph(para)

output_path = "/opt/hermes/profiles/archie/output/Navo_Article_row73_ocean_freight_schedule_reliability.docx"
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print("SAVED:", output_path)
