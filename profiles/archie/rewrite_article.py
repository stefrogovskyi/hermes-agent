
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import re

# Read the HTML content from the file
with open("/opt/hermes/profiles/archie/cache/terminal-output/out-1786800071-605311-deb0.log", "r") as f:
    raw_content = f.read()

# Remove the initial curl output lines (lines 1-4 from the file)
# The actual HTML starts after the line that begins with "0     0    0     0    0     0      0      0 --:--:-- --:--:-- --:--:--     0"
html_start_index = raw_content.find("<!DOCTYPE html>")
if html_start_index != -1:
    html_content = raw_content[html_start_index:]
else:
    html_content = raw_content # Fallback if marker not found

soup = BeautifulSoup(html_content, 'html.parser')

# Extract title
title_tag = soup.find('div', class_='block-name-title')
title = title_tag.h1.text.strip() if title_tag and title_tag.h1 else "No Title Found"

# Extract main content paragraphs and headings
content_div = soup.find('div', class_='blog-single-main-content')
extracted_content = []
if content_div:
    for element in content_div.find_all(['p', 'h2']):
        text = element.get_text(strip=True)
        if text:  # Only add non-empty text
            # Clean up potential inline styles within paragraphs that might have been missed by get_text
            cleaned_text = re.sub(r'style="[^"]*"', '', str(element)).replace('<p>', '').replace('</p>', '').strip()
            extracted_content.append(cleaned_text)

# Reformat and rephrase
rewritten_content = []

# Add the title as a main heading
rewritten_content.append(f"# {title}\n")

# Process each extracted content item for rephrasing or direct inclusion
for item in extracted_content:
    if item.startswith("Every Handover Is a Liability Boundary"):
        rewritten_content.append("\n## Every Transfer Shifts Responsibility\n")
        rewritten_content.append("Every time a shipment moves from one party to another, so does the accountability. The condition of the goods at that precise moment of transfer dictates who is liable for any damage. This is why thorough inspections, detailed reports, and photographic evidence are crucial at every stage, from gate-in to gate-out, including container condition, seal integrity, and cargo stuffing. When these records are comprehensive, resolving a claim is straightforward. However, incomplete documentation often leads to complex negotiations, where the party with better records often has the upper hand, regardless of the actual events.\n")
    elif item.startswith("The Volume Problem Is Not Obvious Until It Is"):
        rewritten_content.append("\n## The Hidden Challenge of High Photo Volumes\n")
        rewritten_content.append("A single shipment can generate numerous photographs: exterior views, seal close-ups, loading sequences, securing methods, damage specifics, delivery conditions, and proof of delivery. For a medium-sized forwarder, this quickly escalates to hundreds of thousands of files annually. Most of these images originate from mobile phones in the field, often with generic filenames, sent via messaging apps, and stored in loosely organized folders. The challenge isn't the existence of the visual data, but whether it can be efficiently retrieved months or years later when a dispute arises.\n")
    elif item.startswith("Finding a Photograph Is the Whole Point"):
        rewritten_content.append("\n## The Essential Role of Retrievable Documentation\n")
        rewritten_content.append("Documentation that cannot be easily found offers no protection. This retrieval issue is a common point of failure in many logistics operations. An image is valuable only when it is accurately linked to a specific booking, container, date, location, and stage of transit. This linkage must be established at the time of capture, not retrospectively. Many organizations explore Digital Asset Management (DAM) platforms, like those outlined in a Cloudinary DAM overview, to learn how metadata, tagging, and search functionalities can ensure files are discoverable by attributes, rather than relying on someone's memory of where they were stored.\n")
    elif item.startswith("Capture in the Field Has to Be Frictionless"):
        rewritten_content.append("\n## Streamlining On-Site Data Capture\n")
        rewritten_content.append("Any system that complicates the capture process will be avoided. Field personnel often operate under time pressure, perhaps on a wet dock with a queue of vehicles. Effective solutions automatically embed context: a mobile app linked to the booking reference, capturing location and timestamp without manual input, and uploading when connectivity allows (or queuing when it doesn't). Asking drivers to rename files or complete forms later is ineffective. The success of a documentation system largely depends on how little it demands from the person taking the photos. While minimal, training is still vital – teaching personnel to include container numbers in the same frame as damage, use objects for scale, and take wide shots before close-ups significantly improves the quality of evidence, making it far more robust than a collection of close-ups without context.\n")
    elif item.startswith("The Same Records Serve More Than Claims"):
        rewritten_content.append("\n## Beyond Claims: The Broader Value of Visual Records\n")
        rewritten_content.append("Well-organized visual documentation provides multiple benefits beyond just claims resolution. Damage photographs, when aggregated, can reveal patterns in recurring problems related to specific shipping lanes, terminals, or packing methods. Detailed condition records strengthen insurance negotiations with concrete evidence. Delivery images help minimize disputes with consignees. Moreover, a comprehensive file compiled during the actual shipment is considerably more compelling to surveyors or insurers than information hastily assembled after a claim has been filed defensively.\n")
    elif item.startswith("Documentation Is Moving Toward Structured Data"):
        rewritten_content.append("\n## The Shift Towards Structured Documentation\n")
        rewritten_content.append("The logistics industry is progressively moving towards structured data. International Maritime Organization (IMO) conventions, such as the Facilitation Convention, now mandate electronic information exchange between ships and ports, signaling a broader transition from paper and scanned images to machine-readable records. Businesses that already maintain their visual documentation in an organized, well-described system are well-prepared for this shift. Conversely, those with photographs scattered across various devices, inboxes, and shared drives will face a more costly transition, as the effort to properly categorize and describe this material will still be required, but at a later, more burdensome stage.\n")
    elif item.startswith("Decide the Standard Before the Dispute"):
        rewritten_content.append("\n## Proactive Documentation Prevents Disputes\n")
        rewritten_content.append("Forwarders and carriers who effectively manage claims typically establish clear protocols in advance. Key questions to address include: What needs to be photographed at each stage, and by whom? How will each image be linked to a specific shipment? Where will the images be stored, who can access them, and for how long will they be retained? Retention policies require careful consideration, as claims and limitation periods in this industry often extend far beyond the intuitive deletion point for old files. Answering these questions proactively is significantly easier than scrambling for answers when a damaged container arrives and no gate-in photograph can be found. Furthermore, agreeing on documentation standards with regular counterparties can expedite dispute resolution, as both sides will accept the evidence without procedural challenges.\n")
    elif item.startswith("A container arrives with a crushed corner. Somewhere between the shipper's yard and the consignee's dock, across a truck, a terminal, a vessel and another terminal, that damage occurred. The commercial question of who bears the cost usually comes down to a much smaller question: at which point in the chain does the first photograph showing that damage appear? In freight, the visual record is not administrative overhead. It is the evidence that settles the claim."):
        rewritten_content.append("Imagine a container arriving with a damaged corner. Determining liability for this damage, which occurred somewhere along its journey from the shipper to the consignee—involving trucks, terminals, and vessels—often hinges on a critical piece of evidence: the initial photograph documenting the damage. In the logistics industry, visual records are not mere formalities; they are definitive proof that resolves claims and assigns responsibility.\n")
    else:
        # For any other paragraphs not explicitly rephrased, add them as is
        rewritten_content.append(item + "\n")

# Remove potential duplicates that might arise from adding both original and rephrased content
final_rewritten_content = []
for paragraph in rewritten_content:
    if paragraph not in final_rewritten_content:
        final_rewritten_content.append(paragraph)

# Create a new Document
document = Document()

# Add the title
document.add_heading(title, level=1)

# Add the rewritten content
for paragraph_text in final_rewritten_content[1:]:  # Skip the first item which is the title already added
    if paragraph_text.startswith("## "):
        document.add_heading(paragraph_text.replace("## ", ""), level=2)
    elif paragraph_text.strip():  # Only add non-empty paragraphs
        document.add_paragraph(paragraph_text.strip())

# Save the document
doc_path = "/opt/hermes/profiles/archie/SeaRates_Blog_Rewritten.docx"
document.save(doc_path)

print(f"Rewritten article saved to {doc_path}")
