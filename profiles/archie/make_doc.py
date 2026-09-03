import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import json
import sys

doc = docx.Document()

# Article Data
title_text = "SeaRates November 2024 Product Update"
meta_title_text = "SeaRates Nov 2024 Update: Tracking, Schedules & TMS"
meta_desc_text = "SeaRates November 2024 release updates ocean, air, rail, parcel tracking, ship schedules, Virtual Office controls, and TMS transport cards."

# H1
h1 = doc.add_heading(level=1)
run_h1 = h1.add_run(title_text)
run_h1.font.size = Pt(20)
run_h1.font.bold = True

# Meta block (Italic 9pt)
meta_p = doc.add_paragraph()
meta_run1 = meta_p.add_run(f"Meta Title: {meta_title_text}\nMeta Description: {meta_desc_text}")
meta_run1.font.size = Pt(9)
meta_run1.font.italic = True

# Spacing
doc.add_paragraph()

# Section 1
p1 = doc.add_paragraph("Monthly platform updates for November 2024 deliver functional improvements across the SeaRates ecosystem based on user feedback.")

h2_1 = doc.add_heading("Tracking and Vessel Schedule Enhancements", level=2)

p2 = doc.add_paragraph("Tracking features on web platforms and web-integrated setups receive a refreshed interface alongside a Book Now button connecting directly to the updated Logistics Explorer. Shipping line logos now display tooltips with carrier names. On the backend, data retrieval logic for extended vessel details was upgraded, and a new predictive ETA calculation formula was implemented.")

p3 = doc.add_paragraph("Coverage across transport modes expanded:")

doc.add_paragraph("Air Cargo Tracking added five carriers (AirMax (Peru), Aloha Air Cargo, Corendon Dutch Airlines, Azerbaijan Airlines, and Air Arabia Abu Dhabi), reaching 437 supported airlines.", style='List Bullet')
doc.add_paragraph("Parcel Tracking added provider Leman, bringing total supported services to 2,417, while improving autodetect logic for API requests.", style='List Bullet')
doc.add_paragraph("Rail Tracking API added a dedicated endpoint to fetch supported rail carriers, introduced the container_size_type field to output container specifications such as \"20' Dry Standard\", refined arrival time calculation logic, and published complete API documentation on the Developer Portal.", style='List Bullet')

p4 = doc.add_paragraph("Ship Schedules broadened carrier query options. Searches by Vessel now support Namsung, SITC, Kambara Kisen, CULines, and Sinokor. Search by Points supports Romocean, while search by Port supports Namsung and Kambara Kisen. For port searches, results surface vessels currently present or recorded within the past 48 hours, alongside ships arriving or departing during that same timeframe.")

# Section 2
h2_2 = doc.add_heading("Virtual Office, TMS, and Interface Adjustments", level=2)

p5 = doc.add_paragraph("Virtual Office dashboard analytics now offer deeper breakdowns. Under the Bookings Overview by Shipping Type chart, selecting More Info reveals detailed metrics covering country, transport mode, shipping type, and route. Toggling between Active Bookings and Requests refreshes linked chart graphics and map displays simultaneously.")

p6 = doc.add_paragraph("The Documents section inside the Bookings tab features a completely redesigned layout. Downloaded files are restricted by default to the booking owner and manager. Other involved parties can access downloaded files through the Show button once review and approval are completed.")

p7 = doc.add_paragraph("Logistics Map and TMS tools now allow users to generate custom thumbnail cover images for transport units. In TMS, selecting Transport Name from the list opens the complete transport card within Logistics Map.")

p8 = doc.add_paragraph("Site customization options now include a Button hover color setting inside Search Filter. The Request an IT Quote form adds informative tooltips across 12 tools, including Freight Index, Air Cargo Tracking, Cargo Wizard, CO2 Calculator, Demurrage and Storage Calculator, and World Sea Ports. Additional tooltips cover SeaRates Mobile App, SeaRates Enterprise, Parcel Tracking, Logistics Map Web access, Web integration, and API.")

docx_path = "/opt/hermes/profiles/archie/SeaRates_November_2024_Product_Update.docx"
doc.save(docx_path)
print(f"Saved DOCX to {docx_path}")

# Upload to Google Drive
parent_folder = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
script = "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py"

cmd = [
    "python3", script, "drive", "upload",
    "--parent", parent_folder,
    "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    docx_path
]

res = subprocess.run(cmd, capture_output=True, text=True)
print("Upload Output:")
print(res.stdout)
if res.stderr:
    print("Upload Errors:", res.stderr)
