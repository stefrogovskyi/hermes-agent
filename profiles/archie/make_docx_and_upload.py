import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import json

doc = docx.Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title H1
title_text = "Practical Guide to International Trade Regulations"
p_title = doc.add_heading(title_text, level=1)

# Meta Title & Meta Description in italic 9pt
meta_title = "Guide to International Trade Regulations"
meta_desc = "Learn how tariffs, customs procedures, trade agreements, and documentation shape international trade compliance and logistics."

p_meta = doc.add_paragraph()
run_meta1 = p_meta.add_run(f"Meta Title: {meta_title}\n")
run_meta1.font.size = Pt(9)
run_meta1.font.italic = True
run_meta1.font.color.rgb = RGBColor(100, 100, 100)

run_meta2 = p_meta.add_run(f"Meta Description: {meta_desc}")
run_meta2.font.size = Pt(9)
run_meta2.font.italic = True
run_meta2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph() # spacing

body_markdown = """Cross-border commerce operates through a grid of regulatory checkpoints, where every cargo container carries both physical goods and a paper trail.

International trade compliance determines how products and services move across borders. These frameworks outline permissible items, handling protocols, transport asset management, and applicable import duties at entry points. They also encompass border clearance routines, bilateral trade deals, and health, safety, or environmental restrictions. Since rules vary widely by jurisdiction, trading partners must review export requirements alongside destination market rules.

### Core Elements of Trade Rules

Four distinct mechanisms structure global regulatory oversight.

Governments charge tariffs and duties on imported or exported items. These fees directly alter landed costs, making accurate calculation necessary when structuring freight pricing. Because tax schedules differ by commodity code and country, checking applicable tariffs prior to dispatch preserves financial projections.

Customs authorities manage border entry to confirm shipments align with national laws. Clearance depends on submitting proper documentation, including commercial invoices, bills of lading, origin certificates, and safety compliance records.

Trade agreements establish preferred terms between participating nations. Frameworks like USMCA (formerly NAFTA) and the European Union single market lower tariff barriers, streamline documentation steps, and simplify freight movement.

Certain items face tight import or export restrictions or absolute prohibitions. Products such as weaponry, hazardous substances, and counterfeit goods remain tightly restricted. Checking regulatory status before booking freight avoids compliance disputes.

### Navigating Regulatory Frameworks

Governments update tariffs, customs workflows, and trade policies frequently in response to shifting economic conditions. Tracking these adjustments requires ongoing monitoring. When regional web restrictions block access, VPN tools secure online activity while enabling access to public databases and regulatory notices, as noted on sites like VPNOverview.com.

Initial research should start with government portals, trade associations, and customs offices. Businesses can also consult freight forwarders or licensed customs brokers specializing in cross-border trade.

Complete documentation must be organized before freight moves. Incomplete or inaccurate paperwork triggers transit delays, financial penalties, or rejected entries. Auditing paperwork with digital tools like Smart Documents helps catch errors prior to official filing.

When regulatory requirements appear complex, partnering with trade specialists provides straightforward guidance. Customs brokers, forwarders, and logistics consultants assist with paperwork, address regulatory queries, and maintain statutory compliance.

### Common Pitfalls

Underestimating tariffs creates unexpected financial burdens. Import duties belong in every initial pricing strategy.

Document mistakes represent a frequent source of transport delays and penalties. Verifying invoice details and certificates prevents administrative friction.

Ignoring local regulations can stop freight at the border. Specific packaging rules or regional product restrictions require early attention during shipping preparation.

Contact SeaRates for logistics assistance."""

for block in body_markdown.split("\n\n"):
    block = block.strip()
    if not block:
        continue
    if block.startswith("### "):
        doc.add_heading(block.replace("### ", ""), level=2)
    elif block.startswith("## "):
        doc.add_heading(block.replace("## ", ""), level=2)
    else:
        doc.add_paragraph(block)

filename = "Practical_Guide_to_International_Trade_Regulations.docx"
doc.save(filename)
print(f"Saved DOCX to {filename}")

# Upload to Google Drive
DRIVE_FOLDER_ID = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
CLI_PATH = "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py"

cmd = [
    "python3", CLI_PATH, "drive", "upload",
    filename,
    "--parent", DRIVE_FOLDER_ID,
    "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
]

res = subprocess.run(cmd, capture_output=True, text=True)
print("UPLOAD STDOUT:", res.stdout)
print("UPLOAD STDERR:", res.stderr)

try:
    upload_info = json.loads(res.stdout)
    file_id = upload_info.get("id")
    web_link = upload_info.get("webViewLink", f"https://docs.google.com/document/d/{file_id}/edit")
    print("SUCCESSFUL UPLOAD!")
    print("File ID:", file_id)
    print("Web Link:", web_link)
    with open("upload_result.json", "w") as f:
        json.dump(upload_info, f, indent=2)
except Exception as e:
    print("Error parsing upload output:", e)

