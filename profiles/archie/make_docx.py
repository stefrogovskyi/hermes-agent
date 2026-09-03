import json
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open("rewrite_final.json", "r", encoding="utf-8") as f:
    rewrite = json.load(f)

doc = docx.Document()

# Document title (Heading 1)
title_text = rewrite["title"]
h1 = doc.add_heading(level=1)
run_h1 = h1.add_run(title_text)
run_h1.font.size = Pt(20)
run_h1.font.bold = True

# Meta title and Meta description in italics 9pt
meta_p = doc.add_paragraph()
meta_title_run = meta_p.add_run(f"Meta Title: {rewrite['meta_title']}\n")
meta_title_run.font.size = Pt(9)
meta_title_run.font.italic = True
meta_title_run.font.color.rgb = RGBColor(100, 100, 100)

meta_desc_run = meta_p.add_run(f"Meta Description: {rewrite['meta_description']}")
meta_desc_run.font.size = Pt(9)
meta_desc_run.font.italic = True
meta_desc_run.font.color.rgb = RGBColor(100, 100, 100)

# Spacer
doc.add_paragraph()

# Body text
body_paragraphs = rewrite["body"].split("\n\n")
for p_text in body_paragraphs:
    p_text = p_text.strip()
    if not p_text:
        continue
    # Add as regular paragraph
    p = doc.add_paragraph()
    run = p.add_run(p_text)
    run.font.size = Pt(11)

filename = "Why_Freight_Rates_Shift_Markets_Tech_and_Global_Routes.docx"
doc.save(filename)
print(f"Created DOCX file: {filename}")
