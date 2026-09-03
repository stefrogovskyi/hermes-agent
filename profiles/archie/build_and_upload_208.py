import os
import re
import subprocess
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = "/opt/hermes/profiles/archie/How_to_Prevent_Freight_Shipping_Delays_Rewritten.docx"

doc = Document()

# Title (H1)
h1 = doc.add_heading("Stop Cargo Delays: Practical Tactics for Reliable Freight", level=1)
h1.style.font.name = 'Arial'
h1.style.font.size = Pt(20)
h1.style.font.bold = True

# Meta block in italic 9pt
p_meta = doc.add_paragraph()
r1 = p_meta.add_run("Meta Title: Avoiding Shipping Bottlenecks: A Freight Delivery Guide\n")
r1.font.name = 'Arial'
r1.font.size = Pt(9)
r1.font.italic = True
r1.font.color.rgb = RGBColor(100, 100, 100)

r2 = p_meta.add_run("Meta Description: Proven methods to avoid freight shipping delays, resolve carrier issues, clear port congestion, and track cargo for dependable deliveries.")
r2.font.name = 'Arial'
r2.font.size = Pt(9)
r2.font.italic = True
r2.font.color.rgb = RGBColor(100, 100, 100)

# Body text
body_markdown = """Late freight shipments throw business operations off balance, creating empty shelves and dissatisfied clients. Steering clear of these obstacles requires three core practices: early preparation, clear dialogue with haulers, and attentive cargo monitoring. Thoughtful workflow tweaks keep your goods moving according to schedule.

## Identifying Frequent Drivers of Shipping Disruptions

Identifying what triggers transit setbacks helps you address complications before they freeze your supply network. Delivery lags generally arise from transport provider troubles, harsh weather events, or terminal bottlenecks.

Hauler difficulties represent a leading reason for stalled freight. Mechanical failures, spot market price shifts, labor deficits, and forgotten pickups can halt transit abruptly. Poor dialogue with your transport provider compounds these headaches.

Documentation mistakes and flawed labeling likewise decelerate cargo. Badly worn or damaged markings leave dock workers guessing, which risks misplacing or misrouting freight. Sturdy industrial tags make sure goods reach their target destination despite rugged handling conditions.

Prevent transport provider and labeling complications with these habits:
- Send precise schedules along with consistent status updates to your hauler.
- Attach rugged, weather-proof tags across all cartons.
- Verify every document before vehicles leave the facility.

## Handling Severe Weather Disruptions

Unpredictable weather patterns derail transit timelines. Downpours, blizzards, storms, and floods bring overland and maritime freight to a standstill. Severe tempests trigger highway and rail line closures, while air carriers ground flights when storms strike.

Uncovered outdoor storage in stormy conditions harms cargo. Moisture causes basic identification tags to bleed or peel, obscuring shipment records. Monitoring weather updates across all transit corridors allows you to warn haulers about prospective slowdowns quickly.

Limit environmental risks through three actions:
- Track weather conditions across starting facilities, transfer points, and final receiving sites.
- Wrap all items in protective outer packaging.
- Build buffer days into freight schedules during stormy months.

## Navigating Ocean Terminal and Port Bottlenecks

Freight frequently sits stranded at overloaded marine terminals longer than planned. Staffing deficits, missing machinery, and heavy cargo volumes build extensive queues. Incomplete documentation or mismatched shipping tags prompt customs detentions that prolong these delays.

Lower terminal bottleneck risks using these methods:
- Transmit documentation and freight details prior to arrival.
- Select port facilities recognized for rapid cargo turnover.
- Apply distinct, legible tags built to survive rough handling.

Transferring freight between different transport providers or conveyance methods adds idle waiting time at freight yards. Mapping journeys early and picking proven shipping channels reduces operational pauses.

## Early Planning and Freight Scheduling

Avoiding transport slowdowns rests on diligent preparation before cargo enters transit lanes. Early foresight helps shipping managers spot vulnerabilities quickly, pick trustworthy logistics partners, and streamline delivery schedules.

## Precise Demand and Freight Volume Projection

Reliable volume projections start with evaluating past order records and sales patterns. Examining previous high-volume seasons and monitoring unexpected demand surges reveals when your busiest and quietest transit phases happen.

Organizing these figures within spreadsheets or a transportation management system (TMS) structures the information cleanly. Such software platforms notify you of incoming volume spikes, giving you space to recalibrate logistics tactics prior to bottlenecks.

Refresh your records constantly and review projections often. Partnering directly with sales and manufacturing departments connects teams around unified figures, lowering rush shipments that run an elevated chance of delay.

## Choosing Dependable Transport Providers

Vetting transport companies demands rigorous assessment. Examine their historic performance for punctual arrival alongside their familiarity with your precise cargo category. Inspect client testimonials, safety track records, and sector evaluations.

Solicit client references or inspect punctuality metrics where available. Professional haulers state clear policies and provide attentive support when disruptions happen. Inspect agreement details carefully, paying attention to liability terms and cargo monitoring support.

Avoid defaulting to the lowest quote available. Unusually low prices frequently translate to sluggish transit or extra transfer points, escalating the likelihood of freight loss or setback. Cultivating lasting partnerships with premier transport providers yields better outcomes.

## Adaptable Dispatch and Delivery Timeframes

Granting adaptable shipping slots simplifies truck booking. Offering flexible pickup and drop-off windows instead of uncompromising deadlines wards off expensive distribution jams.

Transport companies usually show wider capacity when you can dispatch twenty-four hours before or after your primary target date. Such leeway makes it easier to navigate trucker shortages, heavy weather, or regional capacity crunches.

If you rely on timed dock appointments, keep scheduling platforms current and broadcast adjustments right away. Informing haulers about schedule shifts allows them to design better transit routes and set up contingency choices if hiccups occur.

## Refining Documentation and Legal Compliance

Flawless paperwork and strict regulatory adherence are imperative for keeping cargo moving. Mistakes or missing paperwork cause unexpected inspections, border holds, and frustrating delays.

### Accurate Shipping Records

Each outgoing consignment requires verified paperwork:
- Bill of Lading (BOL) paperwork
- Commercial sales invoice
- Itemized packing list
- Certificate of Origin (where mandatory)
- Specialized authorization permits for controlled goods

Review every document thoroughly prior to dispatch. Omissions or typos cause shipments to get impounded at inspection gates.

Ensure paper records mirror physical merchandise exactly. Piece counts, gross weights, and item descriptions must match between paperwork and shipping tags. Utilize clear print or legibly typed paperwork, verifying that all mandatory signatures or authorization stamps are affixed before dispatch.

## International Trade Rules and Customs Clearance

Cross-border freight demands adherence to foreign trade statutes, import taxes, and destination mandates. Investigate local regulations for every target nation prior to shipping to prevent unplanned detentions or surprise penalties.

Customs clearance remains a frequent cause of shipping bottlenecks. Confirm that documentation aligns with the legal standards of the destination territory. Identify restricted goods, confirm import and export licenses, and secure required certificates for restricted merchandise or food products.

Collaborating with a customs clearance broker or freight intermediary streamlines regulatory compliance. These experts audit paperwork to head off penalties, cargo seizures, or hold-ups caused by omissions or clerical mistakes.

## Utilizing Digital Tools for Real-Time Tracking

Modern software provides clear line-of-sight over cargo movements, allowing swift adjustments when obstacles arise.

### Cargo Visibility Platforms

Monitoring systems utilize GPS sensors and barcode scans to transmit live coordinates and updated delivery windows to computers and mobile devices.

Freight directors can monitor interactive maps, graphical reports, and status displays. Select platforms enable staff to inspect photographic evidence or digital paperwork for proof of delivery. Accessing live tracking data eliminates constant phone calls to haulers and removes uncertainty regarding freight progress.

Pinpointing exact shipment coordinates empowers your team to act swiftly when storms, highway bottlenecks, or customs holds create delays. Continuous visibility improves stock management and keeps buyers updated.

### Automated Status Notifications

Instant alerts notify shipping managers about critical milestones along the route. Systems send updates through SMS text messages, email, or application alerts.

Standard notifications cover:
- Dispatch and final delivery receipts
- Unscheduled detours or travel delays
- Customs holds and physical inspection occurrences

Automated notifications prevent missed updates and allow personnel to react swiftly, whether that means dispatching a contingent hauler or notifying a buyer of schedule adjustments. Automated notifications ease operational tension and remove the need for manual status checks.

## Constructing Robust Communication Channels

Open dialogue keeps cargo moving and resolves complications before they spiral. Quick replies and transparent information exchange ward off misunderstandings and operational stalls.

### Partnering With Logistics Providers

Dependable freight transit relies on solid alliances with haulers and transport intermediaries. Establish explicit expectations, responsibilities, and schedules in formal documents. Disclose full shipping specifics ahead of time, including loading dates, destination addresses, cargo weights, and handling notes.

Set up recurring phone or email check-ins to evaluate active transit movements and tackle upcoming friction points. Employ collaborative cloud portals or basic spreadsheets to ensure all parties review matching figures. Store logs of discussions and routing adjustments in a single repository.

### Early Problem Mitigation

Acting promptly when obstacles crop up keeps target dates intact. Configure warnings for operational exceptions, such as transit stalls, cargo damage, or altered routing, and designate a staff member to track goods and intervene.

Execute a four-part strategy whenever disruptions arise:
1. Determine the root cause.
2. Inform your transit provider and recipient.
3. Log corrective steps and progress in writing.
4. Validate that the resolution holds prior to concluding the file.

Keep a current directory of key contacts across all transport partners. Fast access to directory details allows you to connect with necessary personnel without losing time.

## Summary

Averting cargo delays relies on deliberate groundwork and careful attention to process details. Keep paperwork prepared and re-verify shipping records before dispatching freight. Maintain active dialogue with transportation vendors through calls or messages to detect friction points early. Prepare contingency plans, respond promptly when anomalies surface, and keep your staff updated. Although certain transit stalls fall outside your influence, concentrating on practical workflow steps ensures freight moves reliably toward destination points."""

for block in body_markdown.split("\n\n"):
    block = block.strip()
    if not block:
        continue
    if block.startswith("## "):
        h = doc.add_heading(block[3:].strip(), level=2)
        h.style.font.name = 'Arial'
    elif block.startswith("### "):
        h = doc.add_heading(block[4:].strip(), level=3)
        h.style.font.name = 'Arial'
    elif block.startswith("- "):
        for line in block.split("\n"):
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
                p.style.font.name = 'Arial'
    elif re.match(r'^\d+\.\s', block):
        for line in block.split("\n"):
            line = line.strip()
            m = re.match(r'^\d+\.\s*(.*)', line)
            if m:
                p = doc.add_paragraph(m.group(1), style='List Number')
                p.style.font.name = 'Arial'
    else:
        p = doc.add_paragraph(block)
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(11)

doc.save(doc_path)
print(f"Docx saved to {doc_path}, size: {os.path.getsize(doc_path)} bytes")

# Upload via google_api.py
drive_folder = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
api_script = "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py"

cmd = [
    "python3", api_script, "drive", "upload", doc_path,
    "--parent", drive_folder,
    "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
]

res = subprocess.run(cmd, capture_output=True, text=True)
print("Upload output:", res.stdout)
if res.stderr:
    print("Upload stderr:", res.stderr)
