import json
import subprocess
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(json_path, output_docx_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    # Title H1
    h1 = doc.add_heading(level=1)
    run_title = h1.add_run(data['title'])
    run_title.font.size = Pt(20)
    run_title.font.bold = True

    # Meta Title & Description in italic 9pt
    p_meta = doc.add_paragraph()
    r_meta_t = p_meta.add_run(f"Meta Title: {data['meta_title']}\n")
    r_meta_t.font.size = Pt(9)
    r_meta_t.font.italic = True
    r_meta_t.font.color.rgb = RGBColor(100, 100, 100)

    r_meta_d = p_meta.add_run(f"Meta Description: {data['meta_description']}")
    r_meta_d.font.size = Pt(9)
    r_meta_d.font.italic = True
    r_meta_d.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph() # spacing

    # Body Markdown parsing
    body = data['body_markdown']
    lines = body.split('\n')

    for line in lines:
        line_s = line.strip()
        if not line_s:
            continue
        if line_s.startswith('## '):
            h2 = doc.add_heading(level=2)
            r = h2.add_run(line_s[3:].strip())
            r.font.size = Pt(14)
            r.font.bold = True
        elif line_s.startswith('# '):
            h1_sec = doc.add_heading(level=1)
            r = h1_sec.add_run(line_s[2:].strip())
            r.font.size = Pt(16)
            r.font.bold = True
        else:
            p = doc.add_paragraph()
            r = p.add_run(line_s)
            r.font.size = Pt(11)

    doc.save(output_docx_path)
    print(f"Docx created successfully at {output_docx_path}")

def upload_to_drive(docx_path):
    folder_id = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
    cmd = [
        "python3",
        "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py",
        "drive",
        "upload",
        docx_path,
        "--parent", folder_id,
        "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    print("Drive upload output:", res.stdout)
    if res.returncode != 0:
        print("Drive upload error:", res.stderr)
        raise Exception(f"Drive upload failed: {res.stderr}")
    
    upload_data = json.loads(res.stdout)
    file_id = upload_data.get('id')
    web_view_link = upload_data.get('webViewLink') or f"https://docs.google.com/document/d/{file_id}/edit"
    print(f"Uploaded file ID: {file_id}, Link: {web_view_link}")
    return file_id, web_view_link

if __name__ == "__main__":
    docx_file = "/opt/hermes/profiles/archie/Protecting_Modern_Logistics_Through_IT_and_Cybersecurity.docx"
    build_docx("/opt/hermes/profiles/archie/final_article_350.json", docx_file)
    file_id, web_link = upload_to_drive(docx_file)
