from docx import Document
from docx.shared import Pt
import re

base = "/opt/hermes/profiles/archie/workdir/row69"
with open(f"{base}/final_rewrite.txt") as f:
    content = f.read()

lines = content.split('\n')
title = [l for l in lines if l.startswith('TITLE:')][0].split('TITLE:',1)[1].strip()
meta_title = [l for l in lines if l.startswith('META_TITLE:')][0].split('META_TITLE:',1)[1].strip()
meta_desc = [l for l in lines if l.startswith('META_DESCRIPTION:')][0].split('META_DESCRIPTION:',1)[1].strip()

# body = everything after the META_DESCRIPTION line
body_start_idx = None
for i, l in enumerate(lines):
    if l.startswith('META_DESCRIPTION:'):
        body_start_idx = i + 1
        break
body_lines = lines[body_start_idx:]

doc = Document()
doc.add_heading(title, level=1)

p_mt = doc.add_paragraph()
run_mt = p_mt.add_run(f"Meta Title: {meta_title}")
run_mt.italic = True
run_mt.font.size = Pt(9)

p_md = doc.add_paragraph()
run_md = p_md.add_run(f"Meta Description: {meta_desc}")
run_md.italic = True
run_md.font.size = Pt(9)

doc.add_paragraph("")

for line in body_lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    else:
        doc.add_paragraph(line)

out_path = f"{base}/Navo_Article_Row69.docx"
doc.save(out_path)
print("SAVED:", out_path)
