import re
from docx import Document
from docx.shared import Pt

body = open('final_body.md').read()

TITLE = "Road Shipment Tracking Online: A Practical Guide to SeaRates"
META_TITLE = "Track Road Shipments Online with SeaRates Road Tracking"
META_DESC = "See how SeaRates Road Tracking gives real-time road freight visibility, predictive ETA, multi-carrier tracking, white-label pages and TMS/ERP integration."

doc = Document()
doc.add_heading(TITLE, level=1)
p = doc.add_paragraph()
r = p.add_paragraph() if False else None
p = doc.add_paragraph()
run = p.add_run(f"Meta title: {META_TITLE}")
run.italic = True; run.font.size = Pt(9)
p = doc.add_paragraph()
run = p.add_run(f"Meta description: {META_DESC}")
run.italic = True; run.font.size = Pt(9)

for block in body.strip().split('\n\n'):
    b = block.strip()
    if not b:
        continue
    if b.startswith('## '):
        doc.add_heading(b[3:].strip(), level=2)
    elif b.startswith('- '):
        for line in b.split('\n'):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
    else:
        doc.add_paragraph(b.replace('\n', ' '))

doc.save('SeaRates_Road_Tracking_Rewrite_Row117.docx')
print("saved")
