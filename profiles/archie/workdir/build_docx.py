import re
from docx import Document
from docx.shared import Pt

with open('/opt/hermes/profiles/archie/workdir/searates_final_v3.md', encoding='utf-8') as f:
    content = f.read()

title = re.search(r'# TITLE\n(.+)', content).group(1).strip()
meta_title = re.search(r'# META-TITLE\n(.+)', content).group(1).strip()
meta_desc = re.search(r'# META-DESCRIPTION\n(.+)', content).group(1).strip()
body = content.split('# BODY')[1].strip()

doc = Document()

# H1 title
h1 = doc.add_heading(title, level=1)

# Italic 9pt meta fields
p_mt = doc.add_paragraph()
run = p_mt.add_run(f"Meta Title: {meta_title}")
run.italic = True
run.font.size = Pt(9)

p_md = doc.add_paragraph()
run = p_md.add_run(f"Meta Description: {meta_desc}")
run.italic = True
run.font.size = Pt(9)

doc.add_paragraph()  # spacer

# Body: split by lines, detect ## headings vs normal paragraphs
for line in body.split('\n'):
    line = line.strip()
    if not line:
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    else:
        doc.add_paragraph(line)

out_path = '/opt/hermes/profiles/archie/workdir/Smart_Shipping_Means_Reading_Markets_Spreading_Risk.docx'
doc.save(out_path)
print("Saved:", out_path)
