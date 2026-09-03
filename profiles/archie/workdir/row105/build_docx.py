import re
from docx import Document
from docx.shared import Pt

base = "/opt/hermes/profiles/archie/workdir/row105/"
text = open(base+"final_rewrite.txt").read().strip()
title = "B2B Clients Expect Amazon-Level Tracking. Here's the Fix"
meta_title = "Why White Label Shipment Tracking Wins B2B Clients"
meta_desc = "B2B buyer expectations now mirror Amazon. See how an embeddable tracking widget with white label shipment tracking keeps clients on your site."

doc = Document()
doc.add_heading(title, level=1)
p = doc.add_paragraph()
r = p.add_run("Meta title: " + meta_title); r.italic = True; r.font.size = Pt(9)
p = doc.add_paragraph()
r = p.add_run("Meta description: " + meta_desc); r.italic = True; r.font.size = Pt(9)

for block in re.split(r"\n\s*\n", text):
    b = block.strip()
    if not b: continue
    if b.startswith("## "):
        doc.add_heading(b[3:].strip(), level=2)
    else:
        doc.add_paragraph(b)

doc.save(base+"row105.docx")
print("saved")
