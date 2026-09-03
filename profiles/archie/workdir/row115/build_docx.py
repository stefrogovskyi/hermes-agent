from docx import Document
from docx.shared import Pt

src = open('/opt/hermes/profiles/archie/workdir/row115/final.md').read()
lines = src.split('\n')
title = lines[0].replace('TITLE: ', '').strip()
mt = lines[1].replace('META-TITLE: ', '').strip()
md = lines[2].replace('META-DESCRIPTION: ', '').strip()
body = '\n'.join(lines[4:]).strip()

doc = Document()
doc.add_heading(title, level=1)
p = doc.add_paragraph()
r = p.add_run(f"Meta title: {mt}\nMeta description: {md}")
r.italic = True
r.font.size = Pt(9)

for block in body.split('\n\n'):
    b = block.strip()
    if not b:
        continue
    if b.startswith('## '):
        doc.add_heading(b[3:].strip(), level=2)
    else:
        doc.add_paragraph(b)

out = '/opt/hermes/profiles/archie/workdir/row115/standardise-multi-carrier-data-one-stream.docx'
doc.save(out)
print("saved", out)
