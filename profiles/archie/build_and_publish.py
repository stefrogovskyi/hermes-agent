import json
import subprocess
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load article data
with open('/opt/hermes/profiles/archie/article_final.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

title = data['title']
meta_title = data['meta_title']
meta_description = data['meta_description']
body = data['body']

# Create docx
doc = Document()

# Add Title (Heading 1)
h1 = doc.add_heading(title, level=1)

# Add Meta info (italic 9pt)
p_meta = doc.add_paragraph()
run_mt = p_meta.add_run(f"Meta Title: {meta_title}\n")
run_mt.font.size = Pt(9)
run_mt.font.italic = True
run_mt.font.color.rgb = RGBColor(100, 100, 100)

run_md = p_meta.add_run(f"Meta Description: {meta_description}")
run_md.font.size = Pt(9)
run_md.font.italic = True
run_md.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph() # spacing

# Add body paragraphs
for para in body.split('\n\n'):
    para_str = para.strip()
    if not para_str:
        continue
    if para_str.startswith('# '):
        doc.add_heading(para_str[2:].strip(), level=1)
    elif para_str.startswith('## '):
        doc.add_heading(para_str[3:].strip(), level=2)
    elif para_str.startswith('### '):
        doc.add_heading(para_str[4:].strip(), level=3)
    else:
        doc.add_paragraph(para_str)

docx_filename = "SeaRates_Transport_Logistic_2025_Munich_Navo.docx"
docx_path = os.path.join('/opt/hermes/profiles/archie', docx_filename)
doc.save(docx_path)
print(f"Docx saved to {docx_path}")

# Upload to Google Drive
cmd = [
    'python3', '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py',
    'drive', 'upload',
    '--parent', '14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z',
    '--mime-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    docx_path
]

res = subprocess.run(cmd, capture_output=True, text=True)
print("Upload stdout:", res.stdout)
print("Upload stderr:", res.stderr)

if res.returncode != 0:
    print("Upload failed!")
    sys.exit(1)

try:
    upload_res = json.loads(res.stdout)
    file_id = upload_res.get('id')
    web_view_link = upload_res.get('webViewLink', f"https://drive.google.com/file/d/{file_id}/view")
    print("FILE_ID:", file_id)
    print("WEB_VIEW_LINK:", web_view_link)
    
    with open('/opt/hermes/profiles/archie/upload_result.json', 'w', encoding='utf-8') as f:
        json.dump(upload_res, f, ensure_ascii=False, indent=2)
except Exception as e:
    print("Failed to parse upload output:", e)
