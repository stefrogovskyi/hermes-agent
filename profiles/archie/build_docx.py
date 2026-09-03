import json
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    with open("candidate_draft.json", "r") as f:
        data = json.load(f)
        
    doc = docx.Document()
    
    # H1 Title
    title_p = doc.add_heading(data["title"], level=1)
    
    # Meta Title in Italic 9pt
    meta_t_p = doc.add_paragraph()
    run_mt_label = meta_t_p.add_run("Meta Title: ")
    run_mt_label.bold = True
    run_mt_label.font.size = Pt(9)
    run_mt_val = meta_t_p.add_run(data["meta_title"])
    run_mt_val.italic = True
    run_mt_val.font.size = Pt(9)
    
    # Meta Description in Italic 9pt
    meta_d_p = doc.add_paragraph()
    run_md_label = meta_d_p.add_run("Meta Description: ")
    run_md_label.bold = True
    run_md_label.font.size = Pt(9)
    run_md_val = meta_d_p.add_run(data["meta_description"])
    run_md_val.italic = True
    run_md_val.font.size = Pt(9)
    
    doc.add_paragraph() # spacing
    
    # Body markdown parsing
    lines = data["body_markdown"].split("\n")
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
        if line_str.startswith("## "):
            doc.add_heading(line_str[3:].strip(), level=2)
        elif line_str.startswith("* ") or line_str.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            # Parse bold if any
            content = line_str[2:].strip()
            parts = content.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # bold part
                    run.bold = True
        else:
            p = doc.add_paragraph()
            parts = line_str.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
                    
    docx_path = "Colocation_Infrastructure_for_Logistics_and_Shipping_Operations.docx"
    doc.save(docx_path)
    print(f"Saved docx to {docx_path}")
    return docx_path

if __name__ == '__main__':
    create_docx()
