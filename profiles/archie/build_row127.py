from docx import Document
from docx.shared import Pt

rew = open('/opt/hermes/profiles/archie/work_row127_rewrite.md').read()
title = rew.split('TITLE: ')[1].split('\n')[0]
mt    = rew.split('META_TITLE: ')[1].split('\n')[0]
md    = rew.split('META_DESCRIPTION: ')[1].split('\n')[0]
body  = '\n'.join(rew.split('\n')[4:]).strip()

doc = Document()
doc.add_heading(title, level=1)
p = doc.add_paragraph()
r1 = p.add_run(f'Meta title: {mt}\n')
r2 = p.add_run(f'Meta description: {md}')
for r in (r1, r2):
    r.italic = True
    r.font.size = Pt(9)

for block in body.split('\n\n'):
    block = block.strip()
    if not block:
        continue
    if block.startswith('## '):
        doc.add_heading(block[3:], level=2)
    else:
        doc.add_paragraph(block)

out = '/opt/hermes/profiles/archie/row127_dock_worker_injuries.docx'
doc.save(out)
print('saved', out)
