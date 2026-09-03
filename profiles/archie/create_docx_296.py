import json
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('/opt/hermes/profiles/archie/final_data_296.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document()

# Page setup - 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# H1 - Title
h1 = doc.add_heading(data["title"], level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in h1.runs:
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

# Meta Title (Italic 9pt)
p_meta_t = doc.add_paragraph()
run_mt_label = p_meta_t.add_run("Meta Title: ")
run_mt_label.bold = True
run_mt_label.font.size = Pt(9)
run_mt_val = p_meta_t.add_run(data["meta_title"])
run_mt_val.italic = True
run_mt_val.font.size = Pt(9)
run_mt_val.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

# Meta Description (Italic 9pt)
p_meta_d = doc.add_paragraph()
run_md_label = p_meta_d.add_run("Meta Description: ")
run_md_label.bold = True
run_md_label.font.size = Pt(9)
run_md_val = p_meta_d.add_run(data["meta_description"])
run_md_val.italic = True
run_md_val.font.size = Pt(9)
run_md_val.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

# Add spacing line
doc.add_paragraph()

# Body text processing (markdown to docx elements)
body_lines = data["body"].split("\n")
for line in body_lines:
    line_str = line.strip()
    if not line_str:
        continue
    if line_str.startswith("## "):
        h2_text = line_str[3:].strip()
        h2 = doc.add_heading(h2_text, level=2)
        for run in h2.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    else:
        p = doc.add_paragraph()
        run = p.add_run(line_str)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15

docx_path = "/opt/hermes/profiles/archie/Language_Tools_Freight_Software_Communication.docx"
doc.save(docx_path)
print("DOCX created successfully at:", docx_path)
