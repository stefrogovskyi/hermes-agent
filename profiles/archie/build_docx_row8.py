from docx import Document
from docx.shared import Pt

doc = Document()

# H1 title
h1 = doc.add_heading('SeaRates Week 29: Tracking and Rate Updates', level=1)

# Meta title/description italic 9pt
p = doc.add_paragraph()
run = p.add_run('Meta Title: SeaRates Release Notes: Week 29 Updates')
run.italic = True
run.font.size = Pt(9)

p2 = doc.add_paragraph()
run2 = p2.add_run('Meta Description: SeaRates week 29: broader carrier integrations, vessel tracking pricing plans, Load Calculator 3.0 pallets, and multimodal Rail shipments.')
run2.italic = True
run2.font.size = Pt(9)

doc.add_paragraph()

intro = "Small changes pile up fast, and by the time you notice it, the whole platform has shifted under your feet. That's roughly how SeaRates development works: incremental, always shipping something. Check last week's updates first if you haven't already, this one builds on it."
doc.add_paragraph(intro)

doc.add_paragraph("Here's what changed in week 29.")

sections = [
    ("Carrier Integrations: Container and Air Tracking",
     "Container tracking picked up deeper integration work with seven carriers this round: Evergreen, Swire Shipping, Shipping Corporation of India (SCI), CMA CGM, Yang Ming, COSCO, and Shipco Transport. Tracking accuracy and reliability improved for all of them. Air tracking's logistics API connection got stronger too, with better support now in place for Delta Air Lines and T'way Air."),
    ("Vessel Tracking Platform: Pricing Plans and New Languages",
     "The vessel tracking platform now has pricing plans, the main addition to this section this round. Alongside it, the vessel monitoring application picked up interface localization, so it's now available in multiple languages."),
    ("Flight Schedules and Ship Schedules Grow",
     "Airline coverage widened in Flight Schedules with Asiana Airlines and Qantas Airways added to the list. Ship Schedules grew as well, the vessel database gained 9,245 additional fleet records."),
    ("Load Calculator 3.0",
     "Pallets are in. Version 3.0 of the Load Calculator now supports them, including the option to build custom pallet configurations. Load planning gets more flexible, more accurate too."),
    ("Distance & Time: Route Visualization in Color",
     "Route visualization changed this week. Each transportation mode now carries its own dedicated color, making it quicker to scan and follow a route."),
    ("Virtual Office: Facilities Section for Vendors",
     "Virtual Office added a Facilities section. Vendors can now handle their warehouse locations on their own, from setup to day-to-day management."),
    ("Improved Booking Status Visibility",
     "The Booking System got an interface refresh. Booking status synchronization also improved, so what's displayed always matches the actual state of the booking, regardless of what's happening in the browser."),
    ("Rate Management: Alternative LOCODE Lookup",
     "Rate imports, whether through files or API requests, now have an alternative LOCODE lookup option available, making data matching more accurate and imports more reliable."),
    ("Multimodal Rail Shipments in Logistics Explorer",
     "Logistics Explorer added support for multimodal Rail shipments, feeding into freight rate calculations."),
]

for heading, text in sections:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(text)

doc.save('/opt/hermes/profiles/archie/output/SeaRates_Week29_2026.docx')
print("saved")
