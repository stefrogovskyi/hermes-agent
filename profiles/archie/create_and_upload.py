import json
import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set standard margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# H1 Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("STCW Training and Maritime Supply Chain Resilience")
title_run.font.name = "Arial"
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
title_p.paragraph_format.space_after = Pt(12)

# Meta Block (Italic 9pt)
meta_p = doc.add_paragraph()
meta_run1 = meta_p.add_run("Meta Title: ")
meta_run1.font.name = "Arial"
meta_run1.font.size = Pt(9)
meta_run1.font.bold = True
meta_run1.font.italic = True
meta_run1.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

meta_run2 = meta_p.add_run("STCW Safety Standards for Seafarers & Maritime Logistics\n")
meta_run2.font.name = "Arial"
meta_run2.font.size = Pt(9)
meta_run2.font.italic = True
meta_run2.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

meta_run3 = meta_p.add_run("Meta Description: ")
meta_run3.font.name = "Arial"
meta_run3.font.size = Pt(9)
meta_run3.font.bold = True
meta_run3.font.italic = True
meta_run3.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

meta_run4 = meta_p.add_run("FMTC Safety delivers IMO STCW convention compliance through basic safety training, practical drills, and ILT inland courses to protect seafarers and cargo.")
meta_run4.font.name = "Arial"
meta_run4.font.size = Pt(9)
meta_run4.font.italic = True
meta_run4.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

meta_p.paragraph_format.space_after = Pt(18)

body_md = """## Standardized Rules Across Ocean Freight

Thousands of cargo vessels cross international ocean routes every day to move raw commodities and finished products. Ocean freight keeps global markets operating, but ships face harsh conditions where equipment failure or severe storms happen with little warning.

## Core Modules and Emergency Response

The International Maritime Organization sets seafarer competency standards under the Standards of Training, Certification and Watchkeeping regime. Full IMO STCW convention compliance requires crew members to complete certified instruction before taking watch on commercial vessels. FMTC Safety accredited programs cover these requirements through classroom instruction and practical drills.

Four core modules form STCW basic safety training:
- Personal Survival Techniques teaches crew how to launch lifeboats and operate emergency gear.
- Fire Prevention and Firefighting covers fire containment and extinguishing onboard fires effectively.
- Elementary First Aid trains mariners to treat trauma injuries immediately at sea.
- Personal Safety and Social Responsibilities builds teamwork and alertness on board.

Hands-on repetition builds emergency response & survival techniques. When emergency situations occur at sea, crew members must execute safety procedures without delay. Fast action keeps minor incidents from turning into severe losses. SeaRates IT tools support vessel operations by giving managers real-time freight rates visibility and shipment tracking data across ocean routes.

## Inland Waterways and Flexible Class Formats

Maritime work expands beyond deep sea operations into river systems and coastal trade. Specialist ILT programs prepare officers for inland and coastal navigation rules. To keep active crews certified without pulling vessels out of service, FMTC Safety blends online theory modules with practical training at dedicated facilities. Mariners study regulatory coursework remotely before completing physical drills on site.

## Marine Ecosystems and Cargo Protection

Preventing maritime accidents directly strengthens maritime supply chain resilience. Unscheduled vessel delays delay shipments and increase costs for shippers. Trained crews maintain strict equipment checks that reduce shipboard incidents and protect cargo from damage.

Environmental safety depends on crew preparedness. Maritime accidents that cause pollution damage marine ecosystems and harm ocean habitats. Proper training reduces operational risks while meeting international shipping regulations. The SeaRates blog on sustainable shipping details how technological tools and green practices work alongside trained crews to lower environmental impact.

## Industry Guidance Under Michel Hogervorst

FMTC Safety operates under Managing Director Michel Hogervorst, who brings over 15 years of safety training experience across maritime, offshore, wind, and fire sectors. Certified instruction gives seafarers practical skills to maintain safety standards across unpredictable marine environments."""

# Parse body markdown paragraphs and H2 headings
lines = body_md.strip().split('\n')
for line in lines:
    line_str = line.strip()
    if not line_str:
        continue
    if line_str.startswith('## '):
        h2_p = doc.add_paragraph()
        h2_run = h2_p.add_run(line_str[3:].strip())
        h2_run.font.name = "Arial"
        h2_run.font.size = Pt(14)
        h2_run.font.bold = True
        h2_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        h2_p.paragraph_format.space_before = Pt(14)
        h2_p.paragraph_format.space_after = Pt(6)
    elif line_str.startswith('- '):
        bullet_p = doc.add_paragraph(style='List Bullet')
        b_run = bullet_p.add_run(line_str[2:].strip())
        b_run.font.name = "Arial"
        b_run.font.size = Pt(11)
        b_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        bullet_p.paragraph_format.space_after = Pt(4)
    else:
        p = doc.add_paragraph()
        p_run = p.add_run(line_str)
        p_run.font.name = "Arial"
        p_run.font.size = Pt(11)
        p_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        p.paragraph_format.space_after = Pt(8)

docx_filename = "/opt/hermes/profiles/archie/STCW_Training_and_Maritime_Supply_Chain_Resilience.docx"
doc.save(docx_filename)
print(f"Saved DOCX to {docx_filename}")

# Upload via google_api.py
folder_id = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
script_path = "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py"

cmd = [
    "python3", script_path, "drive", "upload", docx_filename,
    "--parent", folder_id,
    "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
]

print("Uploading to Google Drive...")
res = subprocess.run(cmd, capture_output=True, text=True)
print("Upload output:")
print(res.stdout)
if res.stderr:
    print("Upload error:", res.stderr)
