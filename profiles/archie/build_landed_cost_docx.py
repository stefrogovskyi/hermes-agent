#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt

with open('/opt/hermes/profiles/archie/landed_cost_rewrite.md') as f:
    rewrite_full = f.read()

title = re.search(r'^TITLE:\s*(.+)$', rewrite_full, re.M).group(1).strip()
meta_title = re.search(r'^META_TITLE:\s*(.+)$', rewrite_full, re.M).group(1).strip()
meta_desc = re.search(r'^META_DESCRIPTION:\s*(.+)$', rewrite_full, re.M).group(1).strip()
body = rewrite_full.split('---BODY---')[1].strip()

doc = Document()

# H1 title
h1 = doc.add_heading(title, level=1)

# Meta title / description as italic 9pt
p_meta_title = doc.add_paragraph()
run = p_meta_title.add_run(f"Meta Title: {meta_title}")
run.italic = True
run.font.size = Pt(9)

p_meta_desc = doc.add_paragraph()
run2 = p_meta_desc.add_run(f"Meta Description: {meta_desc}")
run2.italic = True
run2.font.size = Pt(9)

doc.add_paragraph()  # spacer

lines = body.split('\n')
for line in lines:
    line = line.rstrip()
    if not line.strip():
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('| '):
        # skip table rows here, handled separately below
        continue
    else:
        doc.add_paragraph(line)

# Handle the table separately: find and insert as a real table
table_match = re.search(r'(\| Cost Component.*?\n(?:\|.*\n)+)', body)
if table_match:
    table_text = table_match.group(1).strip().split('\n')
    rows = [r.strip('|').split('|') for r in table_text if not re.match(r'^\|?\s*-+\s*\|', r)]
    rows = [[c.strip() for c in row] for row in rows]
    if rows:
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = 'Light Grid Accent 1'
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell

doc.save('/opt/hermes/profiles/archie/landed_cost_final.docx')
print("DOCX saved")
