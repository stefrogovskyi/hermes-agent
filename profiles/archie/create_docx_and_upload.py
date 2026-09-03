import re
import json
import subprocess
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(final_file_path, output_docx_path):
    with open(final_file_path, "r", encoding="utf-8") as f:
        text = f.read()

    title_m = re.search(r"^Title:\s*(.*)$", text, re.MULTILINE)
    meta_title_m = re.search(r"^Meta-Title:\s*(.*)$", text, re.MULTILINE)
    meta_desc_m = re.search(r"^Meta-Description:\s*(.*)$", text, re.MULTILINE)

    title = title_m.group(1).strip() if title_m else ""
    meta_title = meta_title_m.group(1).strip() if meta_title_m else ""
    meta_desc = meta_desc_m.group(1).strip() if meta_desc_m else ""

    # Extract body lines
    lines = text.splitlines()
    body_lines = [l for l in lines if not re.match(r"^(Title|Meta-Title|Meta-Description):", l)]

    doc = docx.Document()

    # H1: Title
    h1 = doc.add_heading(title, level=1)
    h1.style.font.size = Pt(20)
    h1.style.font.name = 'Arial'

    # Meta Title (Italic 9pt)
    p_meta1 = doc.add_paragraph()
    run1 = p_meta1.add_run(f"Meta Title: {meta_title}")
    run1.italic = True
    run1.font.size = Pt(9)
    run1.font.name = 'Arial'

    # Meta Description (Italic 9pt)
    p_meta2 = doc.add_paragraph()
    run2 = p_meta2.add_run(f"Meta Description: {meta_desc}")
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = 'Arial'

    # Add space
    doc.add_paragraph("")

    # Parse body lines
    in_bullet_list = False
    for line in body_lines:
        line_s = line.strip()
        if not line_s:
            continue

        if line_s.startswith("### "):
            h2 = doc.add_heading(line_s[4:].strip(), level=2)
            h2.style.font.size = Pt(14)
            h2.style.font.name = 'Arial'
        elif line_s.startswith("#### "):
            h3 = doc.add_heading(line_s[5:].strip(), level=3)
            h3.style.font.size = Pt(12)
            h3.style.font.name = 'Arial'
        elif line_s.startswith("* ") or line_s.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(line_s[2:].strip())
            r.font.size = Pt(11)
            r.font.name = 'Arial'
        else:
            p = doc.add_paragraph()
            r = p.add_run(line_s)
            r.font.size = Pt(11)
            r.font.name = 'Arial'

    doc.save(output_docx_path)
    print(f"Saved DOCX to {output_docx_path}")

def upload_to_drive(docx_path):
    cmd = [
        "python3",
        "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py",
        "drive",
        "upload",
        docx_path,
        "--parent", "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z",
        "--mime-type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        print("Upload failed:", res.stderr)
        return None
    
    print("Upload output raw:", res.stdout)
    try:
        data = json.loads(res.stdout)
        return data
    except Exception as e:
        print("Error parsing upload output JSON:", e)
        return None

if __name__ == "__main__":
    final_txt_path = "/opt/hermes/profiles/archie/final_checked_rewrite.txt"
    docx_path = "/opt/hermes/profiles/archie/SeaRates_Updates_Week_36_2024.docx"
    build_docx(final_txt_path, docx_path)
    res = upload_to_drive(docx_path)
    print("RESULT:", res)
