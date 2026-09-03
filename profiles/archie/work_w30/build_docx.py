from docx import Document
from docx.shared import Pt

title = "SeaRates Week 30: New Carriers, Faster Tracking"
meta_title = "SeaRates Week 30, 2026 Update: New Carriers & Tools"
meta_desc = "New carriers and airlines join SeaRates tracking, plus updates to ship and flight schedules, the logistics map, load calculator, and booking tools."

sections = [
    (None, "Week 30 landed with more names on the carrier list and fewer clicks between a shipment and its status. This week's SeaRates release, put together by Sophia Shkuro, touches container and air tracking, ship and flight schedules, the logistics map, the load calculator, autocomplete, routing, and booking. A couple of these changes are bigger than they look at first glance."),
    ("More Carriers on the Board", "Container tracking picked up ten new shipping line connections this week: TransContainer, Namsung Shipping, Sea Legend Shipping, Dong Young Shipping, Leschaco, Oceanic Star Line, Meratus Line, COSCO Specialized, Atlantic Container Line (ACL), and Jin Jiang Shipping (SHJJ). Each new carrier feeds into the same unified tracking dashboard, adding to a much larger share of global container movement now visible in one place and cutting the need to check a dozen separate carrier portals. More coverage also means the predictive ETA models have more data points to draw on, and exception detection has a better shot at catching a problem before a container sits too long at a terminal racking up demurrage and detention charges.\n\nAir tracking got its own list of additions: My Freighter (Centrum Air), Air Canada, United Airlines, Air New Zealand, SF Airlines, Nippon Cargo Airlines, Air Europa, Qantas, and Uzbekistan Airways. That's nine airlines added to carrier coverage in a single week, a fast pace even by SeaRates' usual release rhythm. For anyone tracking air freight next to ocean freight, this closes gaps that used to force a separate lookup on the airline's own site."),
    ("Schedules Move Too", "Ship Schedules now support Seaboard Marine by Vessel. Wan Hai, KMTC, and Eucon got better support by Points, and Emirates, MTT, and RAL improved by Vessel. On the air side, Flight Schedules added SpiceJet. These are small updates on their own, spread across specific vessels, points, and one added airline schedule."),
    ("The Logistics Map Gets a Table View", "Warehouses finally get a table view. Every tab in the Logistics Map now switches between map and table modes, useful for scanning through dozens of warehouse locations at once. SeaRates also brought back all transport types, Truck, Vessel, Wagon, and Aircraft, in both the Logistics Map and Virtual Office."),
    ("Load Calculator, Autocomplete, and a Few Small Fixes", "The Load Calculator can now import and export Packages and Pallets directly, and users can toggle between metric and imperial units without re-entering every figure by hand. Autocomplete quietly added terminal ports as a filter option too, letting location searches target a specific terminal address directly."),
    ("Smarter Routes and Faster Bookings", "Logistics Explorer changed how multimodal routes get calculated. For LWL shipments, LTL transportation is now automatically added as the first mile of the route, and port search now prioritizes locations in the country where the shipment originates. Both changes produce more relevant routing results without extra manual filtering, useful for anyone tracking a shipment across multiple transport modes.\n\nThe Booking System picked up improvements to booking creation for LWL shipments, moving another step toward real booking automation. The real-time world map used to display shipment routes also got an update this week, so watching a route unfold should feel a little less laggy than before."),
]

doc = Document()

h1 = doc.add_heading(title, level=1)

meta_p = doc.add_paragraph()
run1 = meta_p.add_run(f"Meta title: {meta_title}")
run1.italic = True
run1.font.size = Pt(9)
meta_p.add_run().add_break()
run2 = meta_p.add_run(f"Meta description: {meta_desc}")
run2.italic = True
run2.font.size = Pt(9)

for heading, text in sections:
    if heading:
        doc.add_heading(heading, level=2)
    for para in text.split("\n\n"):
        doc.add_paragraph(para)

doc.save("/opt/hermes/profiles/archie/work_w30/SeaRates_Week_30_2026.docx")
print("saved")
