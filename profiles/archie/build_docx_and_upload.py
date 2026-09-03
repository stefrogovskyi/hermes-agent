import json
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('/opt/hermes/profiles/archie/subagent1_output.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

title = data['title']
meta_title = data['meta_title']
meta_description = data['meta_description']
body = data['body_markdown']

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title H1
h1 = doc.add_heading(level=1)
h1_run = h1.add_run(title)
h1_run.font.name = 'Arial'
h1_run.font.size = Pt(20)
h1_run.font.bold = True
h1_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

# Meta Title (Italic 9pt)
meta_p1 = doc.add_paragraph()
m_title_run = meta_p1.add_run(f"Meta Title: {meta_title}")
m_title_run.font.name = 'Arial'
m_title_run.font.size = Pt(9)
m_title_run.font.italic = True
m_title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Meta Description (Italic 9pt)
meta_p2 = doc.add_paragraph()
m_desc_run = meta_p2.add_run(f"Meta Description: {meta_description}")
m_desc_run.font.name = 'Arial'
m_desc_run.font.size = Pt(9)
m_desc_run.font.italic = True
m_desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph() # Spacer

# Parse Body Markdown
lines = body.split('\n')
for line in lines:
    line_str = line.strip()
    if not line_str:
        continue
    if line_str.startswith('### ') or line_str.startswith('## '):
        h_text = line_str.lstrip('#').strip()
        h2 = doc.add_heading(level=2)
        h2_run = h2.add_run(h_text)
        h2_run.font.name = 'Arial'
        h2_run.font.size = Pt(14)
        h2_run.font.bold = True
        h2_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    else:
        p = doc.add_paragraph()
        p_run = p.add_run(line_str)
        p_run.font.name = 'Arial'
        p_run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)

docx_path = "/opt/hermes/profiles/archie/TILOG_LOGISTIX_2024_SeaRates_Bangkok.docx"
doc.save(docx_path)
print(f"Saved DOCX to {docx_path}")

# Upload to Google Drive folder 14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z
parent_id = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

cmd = [
    'python3',
    '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py',
    'drive', 'upload',
    '--parent', parent_id,
    '--mime-type', mime_type,
    docx_path
]

res = subprocess.run(cmd, capture_output=True, text=True)
print("UPLOAD STDOUT:", res.stdout)
print("UPLOAD STDERR:", res.stderr)

if res.returncode == 0:
    upload_data = json.loads(res.stdout)
    file_id = upload_data.get('id')
    web_view_link = upload_data.get('webViewLink') or f"https://docs.google.com/document/d/{file_id}/edit"
    print(f"SUCCESS: File ID = {file_id}, Link = {web_view_link}")
    with open('/opt/hermes/profiles/archie/upload_result.json', 'w', encoding='utf-8') as f:
        json.dump(upload_data, f, ensure_ascii=False, indent=2)
else:
    print("Upload failed!")
