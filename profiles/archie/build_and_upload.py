import json
import subprocess
import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def make_docx():
    with open("/opt/hermes/profiles/archie/final_article.json", "r", encoding="utf-8") as f:
        data = json.load(f)
        
    doc = Document()
    
    # Title (H1)
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run(data["title"])
    r1.font.size = Pt(20)
    r1.font.bold = True
    
    # Meta Title & Description (Italic, 9pt)
    p_meta = doc.add_paragraph()
    r_mt = p_meta.add_run(f"Meta Title: {data['meta_title']}\n")
    r_mt.font.italic = True
    r_mt.font.size = Pt(9)
    r_mt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    r_md = p_meta.add_run(f"Meta Description: {data['meta_description']}")
    r_md.font.italic = True
    r_md.font.size = Pt(9)
    r_md.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph() # Spacing
    
    # Body text
    body = data["body"]
    lines = body.split("\n")
    
    for line in lines:
        line_s = line.strip()
        if not line_s:
            continue
        if line_s.startswith("## "):
            h2_text = line_s[3:].strip()
            doc.add_heading(h2_text, level=2)
        elif line_s.startswith("# "):
            h1_text = line_s[2:].strip()
            doc.add_heading(h1_text, level=1)
        else:
            doc.add_paragraph(line_s)
            
    docx_path = "/opt/hermes/profiles/archie/How_to_Build_a_Long_Term_Logistics_Career.docx"
    doc.save(docx_path)
    print(f"Saved DOCX to {docx_path}")
    return docx_path

def upload_to_drive(file_path):
    google_api = "/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py"
    folder_id = "14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z"
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    cmd = [
        "python3", google_api,
        "drive", "upload",
        file_path,
        "--parent", folder_id,
        "--mime-type", mime_type
    ]
    
    res = subprocess.run(cmd, capture_output=True, text=True)
    print("Drive upload output:", res.stdout)
    if res.returncode != 0:
        print("Drive upload error:", res.stderr)
        return None
        
    try:
        out_json = json.loads(res.stdout)
        return out_json
    except Exception as e:
        print("Failed to parse drive response:", e)
        return None

if __name__ == "__main__":
    docx_file = make_docx()
    res = upload_to_drive(docx_file)
    print("Upload result:", res)
