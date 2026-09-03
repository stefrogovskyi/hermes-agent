import os
import re
import json
import subprocess
import docx
from docx import Document

# Updated article components after audit
title = "SeaRates Week 40 Platform Updates: Tracking & Geocoding" # 53 chars
meta_title = "SeaRates Week 40 Updates: Carrier Tracking & Geocoding" # 54 chars
meta_desc = "SeaRates Week 40 updates cover expanded carrier tracking, Geocoding API v0.8 scoring, Ship Schedules additions, and mobile app version 1.2 features." # 151 chars

sections = [
    ("Expanded Carrier Tracking Across Air and Sea Networks",
     "Air cargo tracking software capabilities and air cargo tracking API workflows now feature enhanced integration with major international airlines, including Cathay Pacific Airways, Air Canada, Delta Air Lines, Air India, and FedEx Express. On the container and consolidation side, the tracking system has expanded operational support across leasing companies and logistics providers, specifically Shipco Transport, SETH Shipping, and Vanguard Logistics. For ocean carrier integration within Ship Schedules, search by port functionality now supports Dong Young, Culines, and Sinokor. These additions extend tracking coverage across regional and global trade lanes."),

    ("Geocoding API Scoring and Location Routing Enhancements",
     "Logistics geocoding autocomplete receives a functional update in Geocoding API version 0.8. A location scoring algorithm now prioritizes frequently selected hubs, placing most-used origin and destination points at the top of autocomplete queries. Within the Request a Quote workflow, city inputs now benefit from refined spatial logic that identifies nearest commercial ports with higher accuracy. The platform Contact Us form has also been updated to process user inquiries more efficiently."),

    ("Interface Redesigns Across SeaRates and AirRates Web Pages",
     "Visual design and content revisions have rolled out across key web properties within the digital supply chain platform. The SeaRates Affiliate Program page and the Find Freight Routes tool feature revised layouts to clarify partnership structures and route lookup pathways. AirRates.com has updated its primary homepage interface to reflect current service capabilities."),

    ("Release Versions and Platform Deployments",
     "Several platform services received updates and new version deployments during this release cycle. Web users can access the updated Air Cargo Tracking Web Version alongside the Unified Tracking System WEB and the updated Map platform. Development updates include Geocoding API / Autocomplete service Version 0.8, the New Version of Route Planner API, Freight Index 1.0, Booking System Version 1.1, and Load Calculator Version 2.2. Mobile logistics workflows now run on Mobile App Version 1.2, which integrates the Request System feature. Surface freight tracking capabilities extend through the Rail Tracking API and Rail Tracking Web on LandRates.com.")
]

# Read original text for n-gram comparison
with open('/opt/hermes/profiles/archie/orig_332.txt', 'r', encoding='utf-8') as f:
    orig_text = f.read()

# 1. Check Em-Dashes
all_text = title + " " + meta_title + " " + meta_desc + " " + " ".join([s[0] + " " + s[1] for s in sections])
emdash_count = sum(all_text.count(char) for char in ['—', '–', '--'])
print(f"Em-dash count: {emdash_count}")

# 2. Check Meta lengths
print(f"Title length: {len(title)} (max 60) -> {'PASS' if len(title) <= 60 else 'FAIL'}")
print(f"Meta Title length: {len(meta_title)} (max 60) -> {'PASS' if len(meta_title) <= 60 else 'FAIL'}")
print(f"Meta Description length: {len(meta_desc)} (max 155) -> {'PASS' if len(meta_desc) <= 155 else 'FAIL'}")

# 3. Check 6-gram overlaps
def clean_words(t):
    t_clean = re.sub(r'[^\w\s]', ' ', t.lower())
    return [w for w in t_clean.split() if w]

orig_words = clean_words(orig_text)
rewrite_words = clean_words(all_text)

orig_6grams = set()
for i in range(len(orig_words) - 5):
    orig_6grams.add(" ".join(orig_words[i:i+6]))

matches = []
for i in range(len(rewrite_words) - 5):
    gram = " ".join(rewrite_words[i:i+6])
    if gram in orig_6grams:
        matches.append(gram)

print(f"6-gram total matches: {len(matches)}")
for m in set(matches):
    print(" Match:", m)

# 4. Build DOCX
output_dir = "/opt/hermes/profiles/archie/output"
os.makedirs(output_dir, exist_ok=True)
docx_path = os.path.join(output_dir, "Navo_Article_332.docx")

doc = Document()
# Title H1
h1 = doc.add_heading(title, level=1)

# Meta info in italics 9pt
p_meta = doc.add_paragraph()
run_m1 = p_meta.add_run(f"Meta Title: {meta_title}\n")
run_m1.italic = True
run_m1.font.size = docx.shared.Pt(9)
run_m2 = p_meta.add_run(f"Meta Description: {meta_desc}")
run_m2.italic = True
run_m2.font.size = docx.shared.Pt(9)

# Sections H2 + Paragraphs
for s_title, s_text in sections:
    doc.add_heading(s_title, level=2)
    doc.add_paragraph(s_text)

doc.save(docx_path)
print(f"Saved DOCX to: {docx_path}")

# Step 8: Upload to Google Drive
drive_cmd = [
    'python3', '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py',
    'drive', 'upload', docx_path,
    '--parent', '14SwSwwYvop7GLM6R0eDTG5ZLlUTLZr-Z',
    '--mime-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
]
res = subprocess.run(drive_cmd, capture_output=True, text=True)
print("Upload output:", res.stdout, res.stderr)

try:
    upload_res = json.loads(res.stdout)
    file_id = upload_res.get('id')
    web_view_link = upload_res.get('webViewLink', f"https://docs.google.com/document/d/{file_id}/edit")
    print("File ID:", file_id)
    print("WebViewLink:", web_view_link)
except Exception as e:
    print("Error parsing upload output:", e)
    file_id = None
    web_view_link = None

# Step 9: Update Google Sheets row 332
# Columns: D=Статус ("Готово"), E=Название статьи на Наво (title), F=Ссылка на Наво / Файл Наво (web_view_link)
if web_view_link:
    update_d = ['python3', '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py', 'sheets', 'update', '1FI4w3NqDJEzhrqcyJKRWu4sCd57oKSMrfgPEt44b63k', "'Блогпосты Сирейтс'!D332", '--values', '[["Готово"]]']
    update_e = ['python3', '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py', 'sheets', 'update', '1FI4w3NqDJEzhrqcyJKRWu4sCd57oKSMrfgPEt44b63k', "'Блогпосты Сирейтс'!E332", '--values', json.dumps([[title]])]
    update_f = ['python3', '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py', 'sheets', 'update', '1FI4w3NqDJEzhrqcyJKRWu4sCd57oKSMrfgPEt44b63k', "'Блогпосты Сирейтс'!F332", '--values', json.dumps([[web_view_link]])]
    
    subprocess.run(update_d)
    subprocess.run(update_e)
    subprocess.run(update_f)
    
    # Read back to verify
    get_cmd = ['python3', '/opt/hermes/profiles/archie/skills/productivity/google-workspace/scripts/google_api.py', 'sheets', 'get', '1FI4w3NqDJEzhrqcyJKRWu4sCd57oKSMrfgPEt44b63k', "'Блогпосты Сирейтс'!A332:G332"]
    res_get = subprocess.run(get_cmd, capture_output=True, text=True)
    print("Verified Row 332 Data:", res_get.stdout)
