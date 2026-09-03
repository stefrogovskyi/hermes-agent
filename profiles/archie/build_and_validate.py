import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Final polished text
title = "5 Costly Sea Freight Traps and How Ocean Shippers Avoid Them"
meta_title = "5 Costly Sea Freight Traps and How Shippers Avoid Them"
meta_desc = "Learn how to avoid five common sea freight traps, reduce demurrage fees, prevent customs delays, and optimize ocean cargo loads with practical steps."

content_md = """Ocean vessels move eighty percent of global trade volume, providing shippers predictable costs and vast cargo capacity on primary trade routes. Operational friction, however, frequently undermines these advantages. Supply chain teams often fall into execution traps that drive up landed costs, trigger demurrage penalties, and cause delivery delays.

These costly disruptions stem from flawed booking habits, careless paperwork validation, disconnected landside coordination, improper container packing, and reactive crisis management. Managing ocean cargo requires structured planning, precise documentation, and proactive tracking across every leg. Shippers can eliminate these expenses by addressing five operational traps.

## Trap 1: Treating Ocean Freight as a Last-Minute Emergency Fallback

Shippers often turn to ocean shipping only when air transport budgets overflow or road networks face capacity bottlenecks. Treating ocean freight as an emergency parachute forces logistics planners into spot markets days before target vessel departures. This creates immediate operational risk.

Ocean carriers build sailing schedules around long-term commitments. Booking containers late leaves shippers exposed to elevated spot rates, congestion surcharges, and container rollings. When vessels sail full, carriers prioritize contract cargo, pushing spot bookings onto later sailings. A rolled container delays shipments by weeks, disrupting production schedules.

Effective ocean freight risk mitigation requires embedding maritime transport into a rolling supply strategy rather than treating carriers as short-notice backup providers. Ocean logistics demands advance commitment to secure vessel space, equipment availability, and stable pricing.

OPERATIONAL CHECKLIST FOR EARLY OCEAN BOOKINGS:
- Maintain a rolling 60-day to 90-day volume forecast with ocean carriers and forwarders.
- Secure 70 to 80 percent of baseline volume under service contracts while reserving 20 percent for spot allocation.
- Confirm ocean bookings and container pick-up orders 21 days before the Cargo Ready Date (CRD).
- Establish fixed cutoff dates for warehouse loading to ensure containers reach the terminal 48 hours before port gating closes.

An industrial machinery exporter relying on air freight attempted to shift three forty-foot containers to sea freight one week before cargo completion. Without contract allocations, the forwarder accepted inflated spot rates. Space constraints caused two container rollings at origin, delaying delivery by twenty-four days and incurring port storage penalties. A quarterly forecasting cadence helps logistics teams lock in space allocations weeks ahead, which protects margins and keeps delivery schedules intact.

## Trap 2: Documentation Errors and Incorrect HS Codes

International maritime trade relies on precise paperwork. A single character discrepancy across commercial invoices, packing lists, Bills of Lading, or export declarations can bring container movements to a halt. Documentation errors remain a leading cause of port holds, customs fines, and insurance claim rejections.

Customs authorities worldwide utilize automated risk profiling systems that cross-reference entry data against vessel manifests and commercial documents. An incorrect Harmonized System (HS) code, a missing digit, or vague descriptions like "general merchandise" trigger regulatory flags. When customs authorities place a hold on a container, the shipment sits in the terminal while officials request corrected documentation or conduct physical inspections. Storage fees accumulate rapidly, and port authorities pass inspection costs directly to the importer.

Discrepancies between Bills of Lading and commercial invoices can invalidate insurance coverage. If cargo suffers damage, adjusters review all shipping records. Inconsistent cargo descriptions or outdated origin certificates give underwriters grounds to deny claims, leaving shippers to absorb financial losses.

Executing rigorous HS code accuracy & customs compliance requires systematic document verification prior to container gating at origin.

PRE-SHIPMENT DOCUMENTATION AUDIT STEPS:
1. Verify six-digit to ten-digit HS codes against destination tariff schedules for all line items.
2. Cross-check cargo descriptions, quantities, and weights across commercial invoices, packing lists, and draft Bills of Lading.
3. Validate origin certificates, safety data sheets (SDS), and import licenses with licensed customs brokers before container loading.
4. Utilize automated validation tools like Smart Documents to flag data mismatches between commercial and transport documents.

A chemical distributor imported liquid additives using outdated six-digit HS codes that did not match destination tariff classifications. Customs impounded four container tanks for two weeks pending document correction and testing. The hold generated six thousand dollars in terminal storage fees and missed the manufacturing deadline. Implementing a mandatory document audit process prior to issuing final Bill of Lading instructions eliminates administrative holds at destination ports.

## Trap 3: Disconnected Last-Mile Drayage and Warehouse Alignment

Landed ocean transport does not end when a vessel moors at the destination berth. The transition from terminal discharge to inland drayage represents a vulnerable supply chain leg. Failing to coordinate vessel discharge with landside transport leads directly to steep financial penalties.

Ocean carriers and terminal operators grant importers free days (typically three to five calendar days) to pick up discharged containers from the terminal yard. Once free time expires, terminals impose demurrage and port storage fees that scale exponentially for every day the container remains inside the gate. Once the container exits the terminal, carriers grant a separate free time window for returning empty equipment. Delaying warehouse unloading incurs detention and chassis fees.

Poor last-mile drayage coordination happens when shippers treat vessel arrival dates as fixed constants rather than moving targets. Weather, port congestion, and berth shifts alter container availability with little notice. If drayage truckers are not scheduled in advance or warehouses lack bay capacity, containers sit idle in port yards while demurrage fees accumulate.

THE 7-POINT LAST-MILE CONTINGENCY PLAN:
1. Primary and secondary drayage contracts: Secure agreements with multiple drayage providers featuring guaranteed chassis availability and pre-negotiated free time extensions.
2. Off-dock container yard storage: Partner with near-port container yards to haul containers out of the terminal before free time expires when warehouses are full.
3. Advance customs clearance: Pre-lodge customs entries five days prior to vessel arrival to enable instant container release upon discharge.
4. Multiple warehouse receiving locations: Maintain secondary drop locations or cross-dock facilities to handle unexpected volume surges.
5. Flexible trucking appointment slots: Utilize drayage carriers with automated terminal appointment booking systems to secure driver slots immediately upon discharge.
6. Marine cargo insurance transit extensions: Ensure insurance policies cover extended landside storage and container drayage legs without coverage gaps.
7. Schedule buffer days: Incorporate three to four days of operational buffer into delivery promises to absorb landside transport delays without commercial penalties.

A consumer goods importer left destination drayage booking until vessel discharge. Local port congestion limited driver availability, delaying container pickup by six days beyond terminal free time. Demurrage charges totaled over seven hundred dollars per container across ten units, wiping out product margins. Real-time container tracking via terminal APIs enables coordinators to arrange drayage trucks three to five days before vessel arrival, sidestepping port delays and terminal fines.

## Trap 4: Poor Container Packing and Space Inefficiencies

Container packing impacts both transport safety and unit freight economics. Improperly secured cargo, unbalanced weight centers, and unoptimized container volume increase transit damage risks and inflate landed shipping costs per item.

Ocean transport exposes cargo to multi-directional forces, including vessel pitching, rolling, and vibration. When cargo is inadequately braced or weight distribution sits off-center, goods shift inside the container. This movement causes cargo crushing, structural container damage, and road transport accidents during drayage. Unbalanced containers fail SOLAS Verified Gross Mass (VGM) checks at port gates, causing turned-away trucks, re-weighing fees, and missed cutoffs.

Moisture damage presents another major hazard inside steel shipping containers during ocean voyages. Ambient temperature changes cause water vapor inside sealed containers to condense on cold walls and ceilings, dripping onto cargo. This condensation damages electronics, warps wooden products, degrades paper packaging, and causes mold growth on textiles.

Achieving systematic container load optimization involves using 3D packing models and protective packaging standards before loading begins.

CONTAINER LOADING AND CARGO PROTECTION CHECKLIST:
- Run 3D Load Calculator software to map weight distribution, balance axle loads, and maximize internal container volume utilization.
- Calculate and place required high-absorption desiccant bags inside the container based on cargo moisture content and transit climate zones.
- Secure cargo pallets using heavy-duty lashing straps, friction mats, and inflatable dunnage bags in accordance with CTU Code packing guidelines.
- Verify that total gross container weight complies with SOLAS VGM regulations and destination highway weight limits prior to sealing.

An electronics exporter packed containers manually without calculating volume distribution, leaving fifteen percent of cubic space empty while placing heavy power units on one side. The off-center weight balance failed highway drayage inspections at destination, requiring cargo re-rigging and repacking at an off-site warehouse. 3D load planning software balances container weight distribution and maximizes usable volume, so shippers pack more inventory into fewer containers without highway inspection failures.

## Trap 5: Unplanned Crisis Rerouting and Emergency Surcharges

Global shipping lanes face frequent operational disruptions, including port labor strikes, adverse weather events, canal blockades, and carrier blank sailings. When unexpected disruptions close primary shipping routes, unprepared shippers make hasty rerouting decisions that cause severe financial loss.

During major logistics crises, ocean carriers introduce emergency surcharges, spot rate premiums, and rush booking fees. Shippers relying on a single carrier contract find themselves trapped when sailings are canceled or ports skipped. Panicked logistics managers pay inflated spot prices to secure space on alternative vessels, only to discover that destination ports along the alternative route face severe congestion.

Managing emergency disruptions requires adopting a structured freight rate comparison & multi-carrier strategy supported by operational contingency buffers. Maintaining relationships across multiple ocean carrier alliances gives shippers alternative routing channels when primary lines experience delays.

EMERGENCY REROUTING AND COST CONTROL PROTOCOLS:
1. Maintain active contracts with carriers across at least two distinct ocean alliances to preserve alternative vessel access.
2. Utilize digital rate engines like Logistics Explorer to compare live spot rates, transit times, and space availability across multiple carriers during disruptions.
3. Build five to seven operational buffer days into ocean transit schedules on high-risk trade routes to absorb sailing cancellations.
4. Establish a dedicated logistics contingency fund equaling five percent of annual freight spend to cover unavoidable emergency surcharges without freezing operations.

An automotive parts distributor shipped exclusively with one carrier alliance. When severe weather caused port blank sailings, the carrier suspended bookings for three weeks. Forced to secure emergency space with an independent line, the distributor paid double standard freight rates plus panic booking surcharges to meet factory deadlines. Logistics teams that maintain multi-carrier contracts and leverage digital rate tools pivot cargo movements smoothly during disruptions, avoiding extreme spot rate surcharges.

Eliminating these five sea freight traps turns ocean shipping into a reliable supply chain advantage. Shippers who standardize documentation, plan landside drayage early, optimize container space, and maintain flexible carrier strategies protect their profit margins and build resilient global trade operations."""

# Programmatic checks
print("--- STEP 7 PROGRAMMATIC AUDIT ---")

# 1. Em-dash check
full_text = f"{title}\n{meta_title}\n{meta_desc}\n{content_md}"
em_dash_count = full_text.count("—") + full_text.count("--")
print(f"1. Em-dash count: {em_dash_count} (Must be 0)")

# 2. Length checks
print(f"2. Title length: {len(title)} chars (Max 68)")
print(f"   Meta Title length: {len(meta_title)} chars (Max 60)")
print(f"   Meta Description length: {len(meta_desc)} chars (Max 155)")

# 3. Word count
words = re.findall(r'\b\w+\b', content_md)
print(f"3. Word count: {len(words)} words")

# 4. 6-gram overlap with original
orig_text = """Today we're going to play a game called wrong logistics a scenario of your shipment where everything that could go wrong did so. Based on such situations, we will analyze the most common mistakes in maritime cargo transportation and the consequences of making them. At the same time, we won't leave you with only negative impressions but invite you to learn from positive experiences, as we'll consider how to avoid each of these failures. Planning the logistics right is key to making it successful, but what exactly does that mean? What kind of transport should the carrier pick? How should the shipper compare quotes? These and other questions might throw you off track. Let's take a look at some mistakes that can be easily avoided with the right guidance. 1. Sea freight only a fallback option. Our first pitfall would be to choose sea transport only when other options become too expensive or unavailable. From the outset, we will consider sea shipping only as a backup plan in case air transport proves costly. Therefore, we will not plan them, hoping that other transportation options will be more profitable. However, when air transportation started to raise prices because of its high seasonality, sea transportation would be the only one option. Consequences: Restricted route options and carriers offers: There are very few options for line selection and route optimization if you book sea transportation at the last minute. Overloaded lines and extra expenses: Delays, overloaded vessels, and higher transportation expenses result from the growing demand for the few available sea routes. Extra expenses: Ineffective planning may lead to additional costs for accelerated deliveries or unavoidable force majeure events. Sure, that could have been prevented with more thorough preparation. How to avoid it: Maritime transportation as a part of long-term strategies: Don't just use maritime transportation as a backup option. Integrate it into your overall supply strategy. Plan ahead: You'll have a wider carriers selection and routes the earlier you decide to use sea transportation, which will lower costs and prevent delays. Route selection flexibility: Handle unforeseen circumstances more skillfully and guarantee increased dependability when you have a variety of options from sea shipping carriers. 2. Mistakes in documentation. Once the cargo has been shipped by sea, a new problem will arise incorrect documentation. Upon checking the bill of lading, it will become apparent that the HS codes were entered incorrectly, and some certificates were not updated. Consequences: Customs and port delays: Incomplete paperwork can result in customs delays, missed delivery dates, or even cargo seizure, all of which raise expenses. Penalties and fines: Inaccurate documentation leads to fines for breaking port or customs laws, which would make the transportation process even more difficult. Improper insurance: If cargo is damaged or lost, it could be challenging to get payment from the insurance provider if the documentation is inaccurate. How to avoid: Standardize cargo descriptions by describing items in the same way across all paperwork bills of lading, invoices, and customs declarations and making sure that their HS codes correspond. Careful documentation: put in place a system that ensures all required paperwork is checked before cargo is shipped. The names of the goods and the goods code HS code must match, and every document must be filled out accurately. Ask experts: It is best to engage seasoned logisticians or customs brokers or do it with the Smart Documents by yourself to verify and prepare the documentation for important cargo or products that are subject to particular customs regulations. 3. Poor last mile coordination. Due to Murphy's Law or simply from the realistic consequence of not having put the proper logistic arrangements in place, we shall then face yet another problem uncoordinated last mile logistics at the stage of delivery to port and receipt of cargo. Even if the vessel is there on time, are you sure you have already booked a trucking company to pick up cargo at the port and take it to the warehouse? Consequences: Delayed delivery: Cargo on arrival, but not duly coordinated between port and inland transportation, means extra cost of keeping cargo and much less efficiency in the supply chain. Extra storage charges: The longer the delays at the port, the higher the charges for storage, which can become very fast, unexpected, or frightening. Loss of customer loyalty: Late delivery of cargo will hit the company's brand and customer confidence. How to avoid it: Agree on the details in advance: Before the vessel departures, everything about the last mile should be agreed upon the transport, storage space, exact timing for unloading, and receipt of the cargo. Close cooperation: Keep constant contact with all parties of transportation to immediately react to any change or unexpected situation. Ensure customer peace of mind: Provide continuous real-time monitoring of your customers' shipments. Have a backup plan: In case of logistics problems at the last mile, have an alternative pathway to reduce delays. What should your backup plan look like? 1. Alternative transport companies have backup carriers. 2. Additional storage facilities agreements with several warehouses. 3. Backup routes plan alternative delivery routes. 4. Contingency plan for customs procedures customs brokers to speed up the process. 5. Quick replacement of freight transport use of other types of transport. 6. Additional insurance measures cargo insurance to cover additional costs. 7. Planning for buffer time allow extra time for delays. 4. Packaging and loading shortcomings. Now cargo at the port is ready for transportation, but some serious concerns regarding packaging and incorrect placement of different groups of goods have been discovered. No attention had been paid toward securing the cargo in the container, as in maintaining the load toward the center of the container or reserving the position of liquid cargo, which led to an inadequate moisture shield. All these harmed the cargo in transit. Consequences: Cargo damage: Wrong packing, stuffing, or no moisture protection can cause damage to the goods, especially in high humidity or during long sea voyages. Non-effective use of space: Extra empty spaces in the containers increase the cost of transportation and decrease the efficiency of the logistics. Transshipment delays: Improper cargo placement inside containers makes transshipments harder, hence delaying and increasing costs. How to avoid it: Maritime packaging standards: The cargo must be packaged according to international requirements and securely fastened against damage. Moisture protection: Protective materials must be used for moisture-sensitive cargo. Smart load optimization: Unsafe loading and inefficient space utilization can be solved through a Load Calculator, a specialized tool focusing on load optimization into a container or truck. The tool automatically calculates the best possible positioning according to the size, weight, and type of goods, thus allowing the best use of available space within the container. 5. High costs in crises. In the end, all errors, untimely planning, and delays may lead to the last-minute need to alter routes or look for alternative transportation; this already involves very high costs. You may need to pay extra for priority delivery, fines due to unsafe container loading, or even change routes. Consequences: Increasing costs: paying extra for priority services or urgent orders due to delayed payment is already increasing transport costs. Overcharged for emergency bookings: any change in the schedule may generate extra costs for emergency bookings or for other means of transport. Redirection of cargo: When changing the routes or choosing different ports, you will have the additional costs of transshipment and new transport. How to avoid it: Plan for buffer time: It is always better to add some buffer time to shipping schedules so that one does not have to rush into making it an urgent matter due to any delay or a change. Compare multiple contracts: Entering into agreements with several transport companies or lines will allow you to swiftly change carriers or routes with little financial loss. Choose flexibility in freight: Logistics Explorer helps you quickly analyze various options for transportation in real time by comparing available freight rates and transportation options from trusted carriers. This means your transportation plan can be very quickly amended to fit the most cost-effective route and protect you from the risk of very high costs. Make a financial reserve: A reserve fund will help cover unforeseen expenses and ensure financial stability should the transportation process have to be changed."""

def get_ngrams(text, n=6):
    words = re.findall(r'\b[a-z0-9]+\b', text.lower())
    return set(tuple(words[i:i+n]) for i in range(len(words)-n+1))

orig_6grams = get_ngrams(orig_text, 6)
new_6grams = get_ngrams(content_md, 6)
overlap = orig_6grams.intersection(new_6grams)

print(f"4. 6-gram overlap count: {len(overlap)}")
if overlap:
    for gram in list(overlap)[:5]:
        print("   Overlap:", " ".join(gram))

assert em_dash_count == 0, "Em-dash check failed"
assert len(title) <= 68, "Title length failed"
assert len(meta_title) <= 60, "Meta title length failed"
assert len(meta_desc) <= 155, "Meta description length failed"
assert len(overlap) == 0, "6-gram overlap failed"

print("\nALL PROGRAMMATIC CHECKS PASSED SUCCESSFULLY!")

# STEP 8: Assemble DOCX
print("\n--- STEP 8: BUILDING DOCX ---")
doc = Document()

# H1 Title
h1 = doc.add_heading(title, level=1)

# Meta info in italic 9pt
p_meta1 = doc.add_paragraph()
run_m1 = p_meta1.add_run(f"Meta Title: {meta_title}")
run_m1.italic = True
run_m1.font.size = Pt(9)
run_m1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p_meta2 = doc.add_paragraph()
run_m2 = p_meta2.add_run(f"Meta Description: {meta_desc}")
run_m2.italic = True
run_m2.font.size = Pt(9)
run_m2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("") # spacing

# Parse markdown sections into DOCX elements
paragraphs = content_md.split("\n\n")
for block in paragraphs:
    block = block.strip()
    if not block:
        continue
    if block.startswith("## "):
        doc.add_heading(block[3:].strip(), level=2)
    elif block.startswith("### "):
        doc.add_heading(block[4:].strip(), level=3)
    else:
        # Regular paragraph or checklist block
        lines = block.split("\n")
        p = doc.add_paragraph()
        for i, line in enumerate(lines):
            line_str = line.strip()
            if i > 0:
                p = doc.add_paragraph()
            p.add_run(line_str)

docx_path = "/opt/hermes/profiles/archie/output/Navo_Article_173.docx"
doc.save(docx_path)
print(f"Saved DOCX to {docx_path}")
